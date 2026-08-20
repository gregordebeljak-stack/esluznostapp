"""
owner_engine.py
-----------------
Podpora za VEČ LASTNIKOV/LASTNIC iste nepremičnine (npr. zakonca ali drugi
solastniki, ki v pogodbi nastopajo vsak s svojim imenom, naslovom in EMŠO).

Predloge imajo pogosto vnaprej pripravljeni DVA (ali več) "sloti" za vpis
lastnika - npr. ločeni vrstici v tabeli za "lastnico" in "lastnika". Oba
slota v izvirni predlogi vsebujeta rdeč vzorčni tekst v obliki
imena+naslova, zato ju `matching_engine.match_and_apply` (ki en sam slovar
podatkov `extracted` zapiše v VSA ujemajoča polja) privzeto napolni z ISTIM
(prvim) lastnikom. Ta modul po osnovnem klicu `match_and_apply`:

1. Popravi vsak naslednji vnaprej pripravljeni slot z naslednjim lastnikom
   (`assign_owners_to_slots`).
2. Če je lastnikov VEČ kot je v predlogi vnaprej pripravljenih slotov, doda
   nove vrstice po zgledu `multi_parcel_engine.expand_parcel_rows`
   (`expand_owner_rows`).
3. V PROSTEM BESEDILU (izven tabel) - kjer se lastnika omenjata zgolj po
   imenu (npr. "... katera lastnika sta služnostna zavezanca: Janez Novak,
   Ana Novak") - nadomesti posamezno ime s SEZNAMOM VSEH lastnikov, ločenim
   z vejico (`expand_free_text_owner_mentions`).
4. Zapolni polja, ki so v predlogi ŽE VNAPREJ zapisana kot SEZNAM VEČ imen,
   ločenih z vejico (npr. "Ana Šabeder, Silvester Šabeder" - tipično v
   povzetni tabeli s parcelami/lastniki) - `matching_engine` takega polja
   sploh ne zazna (ni niti golo ime niti ime+naslov), zato ga
   `match_and_apply` ne dotakne - to ureja `fill_owner_list_summary_fields`.

Prvi lastnik (owners[0]) je v predlogi že pravilno zapisan prek običajnega
matching_engine.match_and_apply - ta modul poskrbi za DRUGEGA in naslednje.

Opomba: če se slota za lastnika in slota za parcelo nahajata v ISTI tabeli,
lahko vrstni red klicanja expand_*_rows funkcij (ta modul vs.
multi_parcel_engine) vpliva na indekse vrstic - v praksi so pri realnih
predlogah tabela lastnikov in tabela parcel ločeni, zato to ni težava.
"""

from __future__ import annotations

import copy
import re
from collections import OrderedDict
from typing import Dict, List, Optional

from docx.shared import RGBColor
from docx.table import _Row

from docx_engine import Field, FieldLocation, apply_field_value
from matching_engine import _NAME_RE, _ADDRESS_RE, _BARE_NAME_RE, _preserve_surname_case

BLACK = RGBColor(0x00, 0x00, 0x00)

# Vloge, ki jih obravnavamo kot "mesto za vpis lastnika" - glej
# matching_engine.classify_role / MatchResult.role.
OWNER_ROLES = {"ime_priimek", "naslov", "ime_priimek+naslov"}

# Manjši seznam znanih izjem - slovenska moška imena, ki se (za razliko od
# večine moških imen) končajo na "a". Preostala imena, ki se končajo na
# "a", štejemo za ženska (Marija, Ana, Petra ...). Hevristika ni 100 %
# zanesljiva za vsa možna imena, a zadostuje za razločevanje med spoloma
# solastnikov na podpisnih mestih predloge (npr. "Služnostna zavezanka" /
# "Služnostni zavezanec").
_MALE_NAME_EXCEPTIONS = {"luka", "jaka", "saša", "nikola", "kosta"}


def _detect_gender(ime_priimek: str) -> Optional[str]:
    """Vrne "Z" (žensko) ali "M" (moško) glede na prvo besedo (ime) v
    `ime_priimek`, ali None, če imena ni mogoče prepoznati (prazen niz)."""
    parts = (ime_priimek or "").strip().split()
    if not parts:
        return None
    name = parts[0].lower()
    if not name:
        return None
    if name in _MALE_NAME_EXCEPTIONS:
        return "M"
    return "Z" if name.endswith("a") else "M"


def _set_location_value(loc: FieldLocation, value: str):
    """Zapiše vrednost v runs ENE KONKRETNE lokacije polja (ne vseh pojavitev
    polja hkrati) in jo obarva črno - enako kot
    multi_parcel_engine._set_location_value, a lokalna kopija, da modul
    ostane samostojen."""
    if not loc.runs:
        return
    loc.runs[0].text = value
    try:
        loc.runs[0].font.color.rgb = BLACK
    except Exception:
        pass
    for extra_run in loc.runs[1:]:
        extra_run.text = ""
        try:
            extra_run.font.color.rgb = BLACK
        except Exception:
            pass


def _set_cell_first_run(cell, value: str):
    """Zapiše vrednost v prvi run prve neprazne odstavka celice (ohrani
    obliko tega run-a), ostale run-e v istem odstavku počisti."""
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = value
            try:
                p.runs[0].font.color.rgb = BLACK
            except Exception:
                pass
            for extra in p.runs[1:]:
                extra.text = ""
            return
    if cell.paragraphs:
        run = cell.paragraphs[0].add_run(value)
        run.font.color.rgb = BLACK


def _format_owner_value(role: str, current_field_value: str, owner: dict) -> str:
    """Zgradi novo vrednost za DANO vlogo/lokacijo na podlagi trenutne
    (dejansko že zapisane, torej vrednosti PRVEGA lastnika) besedilne oblike
    polja - iz nje po potrebi izluščimo morebiten "rep" (npr. ", EMŠO: "), ki
    sledi imenu+naslovu v izvirni predlogi, in ga ohranimo tudi za naslednje
    lastnike."""
    if role == "ime_priimek+naslov":
        trailing = ""
        m = _NAME_RE.match(current_field_value)
        if m:
            rest = m.group("rest")
            addr_m = _ADDRESS_RE.search(rest)
            if addr_m:
                trailing = rest[addr_m.end():]
        return f"{owner.get('ime_priimek', '')}, {owner.get('naslov', '')}{trailing}"
    if role == "ime_priimek":
        return owner.get("ime_priimek", "")
    if role == "naslov":
        return owner.get("naslov", "")
    return current_field_value


def identify_owner_table_locations(fields: Dict[str, Field], match_report) -> List[dict]:
    """Vrne urejen seznam FIZIČNIH pojavitev (ne polj!) "slotov" za lastnika,
    razvrščen po vrstnem redu v dokumentu (tabela, vrstica, stolpec; nato
    prostobesedilna mesta).

    Zajame:
    1. Vse tabelske pojavitve vlog lastnika (kot doslej).
    2. PROSTOBESEDILNA mesta (izven tabel) z vlogo "ime_priimek+naslov" -
       to so npr. podpisna mesta v prostem besedilu ("... Janez Novak,
       Naslov ... EMŠO: ___"), ki so prav tako fizični "sloti" za
       posameznega lastnika, in jih je treba razporediti med lastnike enako
       kot tabelske slote (glej `assign_owners_to_slots`). Golih omemb
       SAMEGA imena (vloga "ime_priimek") v prostem besedilu tu namenoma NE
       zajamemo - te obravnava `expand_free_text_owner_mentions` (nadomesti
       jih s seznamom vseh imen, ločenim z vejico), saj gre za drugačen
       vzorec (ena omemba, ne zaporedje ločenih slotov).

    POMEMBNO: če sta bila v izvirni predlogi dva slota (npr. za lastnico in
    lastnika) zapisana z BESEDILNO IDENTIČNIM vzorčnim imenom+naslovom, ju
    je `docx_engine.parse_docx_fields` združil v EN sam Field z dvema
    lokacijama (`field.locations`) - zato tu namenoma iteriramo po vseh
    lokacijah vsakega ujemajočega polja, ne le po enem vnosu na polje, da
    vsak fizičen slot v dokumentu dobi svoj zaporedni indeks."""
    entries: List[dict] = []
    for order, r in enumerate(match_report):
        if r.role not in OWNER_ROLES:
            continue
        f = fields.get(r.field_id)
        if f is None:
            continue
        for loc in f.locations:
            in_table = loc.table_index is not None and loc.row_index is not None
            if not in_table and r.role != "ime_priimek+naslov":
                continue
            entries.append({
                "field": f,
                "field_id": r.field_id,
                "loc": loc,
                "role": r.role,
                "table_index": loc.table_index,
                "row_index": loc.row_index,
                "col_index": loc.col_index or 0,
                "order": order,
            })
    entries.sort(key=lambda e: (
        0 if e["table_index"] is not None else 1,
        e["table_index"] if e["table_index"] is not None else 0,
        e["row_index"] if e["row_index"] is not None else 0,
        e["col_index"],
        e["order"],
    ))
    return entries


def assign_owners_to_slots(fields: Dict[str, Field], match_report, owners: List[dict]) -> int:
    """Popravi vsak naslednji (2., 3., ...) vnaprej pripravljeni slot za
    lastnika z ustreznim naslednjim lastnikom iz `owners`. Prvi slot
    (owners[0]) je v predlogi že pravilno zapisan prek match_and_apply, zato
    ga tu (praviloma) preskočimo. Če je lastnikov VEČ kot slotov, presežne
    lastnike prevzame `expand_owner_rows` (doda nove vrstice) - tu se za
    zadnji slot uporabi zadnji razpoložljivi lastnik.

    Razporeditev poteka NEODVISNO za vsako FIZIČNO SKUPINO slotov (ista
    tabela + isti stolpec, oz. prosto besedilo) posebej, ne po enem SKUPNEM
    naraščajočem indeksu čez cel dokument. Dokument namreč lahko vsebuje VEČ
    ločenih skupin slotov za lastnika (npr. prostobesedilni sloti v 4. členu
    IN ločeno tabela "Služnostni zavezanec/-ka") - če bi si delile en sam
    naraščajoč indeks, bi poznejši skupini "zmanjkalo" lastnikov in bi se
    vsi njeni sloti napačno napolnili z ZADNJIM lastnikom (owners[-1]),
    namesto da vsaka skupina spet začne pri lastniku[0].

    POMEMBNO: skupine NAMENOMA niso določene po field_id (kot bi se morda
    zdelo naravno), temveč po fizični lokaciji (tabela + stolpec, oz. prosto
    besedilo). V realnih predlogah namreč vsak vnaprej pripravljeni slot
    praviloma vsebuje SVOJE (različno) vzorčno ime - npr. "Ana Šabeder, Na
    griču 12 ..." v prvi vrstici in "Silvester Šabeder, Na griču 12 ..." v
    drugi - zato jih `docx_engine.parse_docx_fields` NE združi v en Field
    (dedup je po točnem besedilu), ampak ustvari dva ločena Field objekta,
    vsak s toliko lokacijami. Če bi grupirali po field_id, bi vsak tak slot
    pristal v svoji lastni skupini velikosti 1 in bi ga spodnja zanka
    (`len(field_entries) <= 1: continue`) preskočila - s tem bi 2., 3., ...
    slot ostal nepopravljen (in bi obdržal vrednost prvega lastnika, ki jo
    je tja nekritično zapisal `match_and_apply`, saj ta isti `extracted`
    slovar zapiše v VSA ujemajoča polja, ne glede na to, koliko slotov je).

    Za polja z GOLIM imenom (vloga "ime_priimek", brez naslova) z več kot
    enim slotom, kadar so med lastniki MEŠANI SPOLI (npr. podpisni mesti
    "Služnostna zavezanka" / "Služnostni zavezanec"), razvrstimo lastnike
    tako, da ženska pride na prvi (v dokumentu zgodnejši), moški pa na
    naslednji slot - to ustreza vrstnemu redu nazivov v predlogi (zavezanka
    pred zavezancem). Če spola niso mešani (npr. dva solastnika istega
    spola), ostane vrstni red nespremenjen.

    Vrne število dejansko popravljenih slotov."""
    if len(owners) <= 1:
        return 0
    entries = identify_owner_table_locations(fields, match_report)
    if len(entries) <= 1:
        return 0

    def _group_key(entry: dict):
        # Fizična skupina: znotraj tabele grupiramo po (tabela, stolpec) -
        # tako vrstice z RAZLIČNIM vzorčnim besedilom (torej različnimi
        # Field objekti), ki pa so del istega "slot-stolpca" (npr. stolpec z
        # imenom+naslovom lastnika), pristanejo v isti skupini. Zunaj tabele
        # (prosto besedilo, npr. podpisna mesta v 4. členu) grupiramo samo
        # po vlogi - to je ločena skupina od katerekoli tabele.
        if entry["table_index"] is not None:
            return ("table", entry["table_index"], entry["col_index"], entry["role"])
        return ("freetext", entry["role"])

    groups: "OrderedDict[tuple, list]" = OrderedDict()
    for entry in entries:
        groups.setdefault(_group_key(entry), []).append(entry)

    changed = 0
    for field_entries in groups.values():
        if len(field_entries) <= 1:
            continue

        role = field_entries[0]["role"]
        owner_order = owners
        if role == "ime_priimek":
            genders = [_detect_gender(o.get("ime_priimek", "")) for o in owners]
            if "Z" in genders and "M" in genders:
                owner_order = sorted(
                    owners,
                    key=lambda o: 0 if _detect_gender(o.get("ime_priimek", "")) == "Z" else 1,
                )

        for i, entry in enumerate(field_entries):
            target = owner_order[i] if i < len(owner_order) else owner_order[-1]
            if i == 0 and target is owners[0]:
                continue  # to mesto je že pravilno napolnjeno prek match_and_apply
            new_val = _format_owner_value(entry["role"], entry["field"].value, target)
            _set_location_value(entry["loc"], new_val)
            changed += 1
    return changed


def expand_owner_rows(doc, fields: Dict[str, Field], match_report, owners: List[dict]) -> int:
    """Če je lastnikov VEČ kot je v predlogi vnaprej pripravljenih slotov
    (vrstic) zanje, podvoji ZADNJO tako vrstico (skupaj z morebitno
    spremljajočo vrstico takoj za njo - npr. "(v nadaljevanju: ...)" - če
    obstaja, da nova vrstica ohrani identično obliko) za vsakega dodatnega
    lastnika in jo zapolni z njegovimi podatki. Vrne število dejansko
    dodanih vrstic (samo za lastnika - vrstice, dodane za dodatne parcele,
    ureja `multi_parcel_engine.expand_parcel_rows`)."""
    if len(owners) <= 1:
        return 0
    # Samo TABELSKI sloti - v prostem besedilu (npr. 4. člen) ni mogoče
    # zanesljivo vstaviti novega odstavka po zgledu obstoječega, zato tam
    # dodajanje novih slotov ni podprto (le razporeditev obstoječih prek
    # assign_owners_to_slots).
    entries = [e for e in identify_owner_table_locations(fields, match_report) if e["table_index"] is not None]
    if not entries or len(owners) <= len(entries):
        return 0  # dovolj vnaprej pripravljenih slotov - te samo popravi assign_owners_to_slots

    last = entries[-1]
    t_idx, r_idx = last["table_index"], last["row_index"]
    if t_idx >= len(doc.tables):
        return 0
    table = doc.tables[t_idx]
    if r_idx >= len(table.rows):
        return 0

    template_tr = table.rows[r_idx]._tr
    companion_tr = table.rows[r_idx + 1]._tr if r_idx + 1 < len(table.rows) else None

    extra_owners = owners[len(entries):]
    anchor_tr = companion_tr if companion_tr is not None else template_tr
    added = 0
    for owner in extra_owners:
        new_tr = copy.deepcopy(template_tr)
        anchor_tr.addnext(new_tr)
        anchor_tr = new_tr
        new_row = _Row(new_tr, table)
        new_val = _format_owner_value(last["role"], last["field"].value, owner)
        for cell in new_row.cells:
            if cell.text.strip():
                _set_cell_first_run(cell, new_val)
                break

        if companion_tr is not None:
            new_companion_tr = copy.deepcopy(companion_tr)
            anchor_tr.addnext(new_companion_tr)
            anchor_tr = new_companion_tr

        added += 1

    return added


def expand_free_text_owner_mentions(fields: Dict[str, Field], match_report, owners: List[dict]) -> int:
    """Ko je lastnikov VEČ KOT EN, poišče omembe SAMEGA IMENA lastnika (brez
    naslova) v PROSTEM BESEDILU predloge (izven tabel - npr. stavek "...
    katera lastnika sta služnostna zavezanca: Janez Novak") in jih nadomesti
    s SEZNAMOM VSEH lastnikov, ločenim z vejico. Polja, ki se pojavijo
    IZKLJUČNO v tabelah (posamezni sloti lastnika), tu niso obravnavana -
    zanje poskrbita `assign_owners_to_slots` in `expand_owner_rows`.

    Vrne število dejansko posodobljenih omemb."""
    if len(owners) <= 1:
        return 0

    names: List[str] = []
    seen = set()
    for o in owners:
        name = (o.get("ime_priimek") or "").strip()
        if name and name not in seen:
            seen.add(name)
            names.append(name)
    if len(names) <= 1:
        return 0

    new_text = ", ".join(names)
    updated = 0
    for r in match_report:
        if r.role != "ime_priimek":
            continue
        f = fields.get(r.field_id)
        if f is None:
            continue
        free_text_locs = [loc for loc in f.locations if loc.table_index is None]
        if not free_text_locs:
            continue  # polje je samo v tabeli - to ureja assign_owners_to_slots/expand_owner_rows
        for loc in free_text_locs:
            _set_location_value(loc, new_text)
        updated += 1

    return updated


def dedupe_owners(records: List[dict]) -> List[dict]:
    """Iz seznama zapisov (parcels_data - lahko vsebuje več vrstic za istega
    lastnika, npr. če ima več parcel) izlušči SEZNAM UNIKATNIH lastnikov,
    razvrščenih po vrstnem redu prvega pojava, prepoznanih po paru
    (ime_priimek, naslov)."""
    seen = set()
    owners: List[dict] = []
    for rec in records:
        key = (
            (rec.get("ime_priimek") or "").strip().lower(),
            (rec.get("naslov") or "").strip().lower(),
        )
        if key not in seen:
            seen.add(key)
            owners.append(rec)
    return owners


def fill_owner_list_summary_fields(fields: Dict[str, Field], match_report, owners: List[dict]) -> int:
    """Zapolni polja, katerih IZVIRNO (predlogino) besedilo je SEZNAM VEČ
    imen, ločenih z vejico (npr. "Ana Šabeder, Silvester Šabeder") - to so
    povzetna mesta (tipično v tabeli, ki povzame parcelo/parcele SKUPAJ z
    vsemi njenimi lastniki, npr. stolpec "parc. št. | k.o. | delež" z
    naslovno vrstico "Ime Priimek1, Ime Priimek2"), ki jih
    `matching_engine._build_replacement` NE zazna - prepozna namreč samo
    GOLO ime (ena oseba) ali ime+naslov, ne pa seznama VEČ imen - zato jih
    `match_and_apply` sploh ne vključi v `match_report` in ostanejo
    nedotaknjena (rdeča, z vzorčnimi imeni iz predloge).

    Prepoznavanje: polje je "seznam imen", če se njegova vrednost pri
    razdelitvi po vejici razdeli na VSAJ 2 dela in VSAK del sam zase izgleda
    kot osebno ime (enak vzorec kot `matching_engine._BARE_NAME_RE` - "Ime
    Priimek", 2-4 besede, vsaka z veliko začetnico). Namenoma preskočimo
    polja, ki jih je `match_and_apply` že obravnaval (so v `match_report`),
    da po nepotrebnem ne bi podvojili/prepisali polj, za katera skrbijo
    `assign_owners_to_slots` / `expand_free_text_owner_mentions`.

    Če je zaznanih pravih lastnikov (owners) MANJ kot 2, take zamenjave ni
    mogoče smiselno narediti (seznam bi bil dolg 1 ime) - v tem primeru
    funkcija ne naredi ničesar in vrne 0.

    Vrne število dejansko posodobljenih POLJ (za informativno sporočilo v
    UI)."""
    names: List[str] = []
    seen = set()
    for o in owners:
        name = (o.get("ime_priimek") or "").strip()
        if name and name not in seen:
            seen.add(name)
            names.append(name)
    if len(names) < 2:
        return 0

    joined = ", ".join(names)
    matched_field_ids = {r.field_id for r in match_report}

    updated = 0
    for fid, f in fields.items():
        if f.is_blank_placeholder or fid in matched_field_ids:
            continue
        parts = [p.strip() for p in f.value.split(",")]
        if len(parts) < 2 or not all(_BARE_NAME_RE.match(p) for p in parts):
            continue
        apply_field_value(f, joined)
        updated += 1

    return updated
