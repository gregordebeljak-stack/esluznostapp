"""
multi_parcel_engine.py
-----------------------
Podpora za VEČ parcel istega lastnika (npr. lastnik/lastnica/občina ima v
lasti več parcel in za vsako obstaja svoj PDF izpisek iz zemljiške knjige).

Ko je naloženih več PDF-jev, to modul:
1. Ugotovi, katere vrstice v tabelah .docx predloge vsebujejo "čisto"
   parcelno-specifične podatke (parcelna_stevilka / katastrska_obcina / delez)
   - to so vrstice, ki jih je treba podvojiti za vsako dodatno parcelo.
2. Podvoji te vrstice (ohranjajoč IDENTIČNO obliko/pisave, saj gre za dobeseden
   XML kopiranje) in jih zapolni s podatki naslednjih parcel.
3. V PROSTEM BESEDILU (izven tabel) - npr. v stavku "...da se na zemljišču
   s parc. št.: X, k.o. Y izrecno in nepogojno dovoljuje..." - nadomesti
   posamezno parcelno številko (in po potrebi katastrsko občino, če se med
   parcelami razlikuje) s SEZNAMOM VSEH parcel lastnika, kot so navedene v
   tabelah - to ureja `expand_free_text_parcel_mentions`.

Prva parcela (parcels[0]) je v predlogi že zapisana prek običajnega
matching_engine.match_and_apply - ta modul doda VRSTICE ZA PREOSTALE parcele
(v tabelah) oz. razširi omembe v prostem besedilu.
"""

from __future__ import annotations

import copy
import re
from typing import Dict, List, Tuple

from docx.shared import RGBColor
from docx.table import _Row

from docx_engine import Field, apply_field_value

BLACK = RGBColor(0x00, 0x00, 0x00)
PARCEL_ROW_ROLES = {"parcelna_stevilka", "katastrska_obcina", "delez"}
_FRACTION_RE = re.compile(r"(\d+\s*/\s*\d+)")


def _group_parcel_locations_by_cell(fields: Dict[str, Field], match_report) -> Dict[Tuple[int, int, int, str], List[Tuple[str, Field, object]]]:
    """Združi VSE (vloga, polje, lokacija) trojice z vlogo parcelna_stevilka/
    katastrska_obcina/delez po FIZIČNI celici tabele (table_index, row_index,
    col_index) - ne glede na to, ali gre za isto ali RAZLIČNA Field polja.

    To je pomembno, ker predloga v isti celici pogosto vnaprej pripravi VEČ
    ločenih odstavkov z RAZLIČNIM vzorčnim besedilom (npr. '949/22' in
    '949/23' - dva možna zapisa za do 2 različni parcelni številki) -
    docx_engine.parse_docx_fields ju zato NE združi v eno Field polje (dedup
    je po točnem besedilu), ampak ustvari dve ločeni polji, vsako s svojo
    lokacijo v ISTI celici. Funkcije, ki gledajo samo `field.locations`
    enega polja naenkrat (kot je bilo prej), take medsebojne podvojitve v
    isti celici ne opazijo."""
    by_cell: Dict[Tuple[int, int, int, str], List[Tuple[str, Field, object]]] = {}
    for r in match_report:
        if r.role not in PARCEL_ROW_ROLES:
            continue
        f = fields.get(r.field_id)
        if f is None:
            continue
        for loc in f.locations:
            if loc.table_index is None or loc.row_index is None:
                continue
            key = (loc.table_index, loc.row_index, loc.col_index or 0, r.role)
            by_cell.setdefault(key, []).append((r.role, f, loc))
    return by_cell


def identify_parcel_row_groups(fields: Dict[str, Field], match_report) -> Dict[Tuple[int, int], Dict[int, str]]:
    """Vrne {(table_index, row_index): {col_index: vloga}} za vse vrstice v
    tabelah, ki vsebujejo vsaj enega od treh "čisto" parcelnih podatkov.
    Te vrstice je treba podvojiti, če je parcel več kot ena."""
    groups: Dict[Tuple[int, int], Dict[int, str]] = {}
    for r in match_report:
        if r.role not in PARCEL_ROW_ROLES:
            continue
        f = fields.get(r.field_id)
        if f is None:
            continue
        for loc in f.locations:
            if loc.table_index is None or loc.row_index is None or loc.col_index is None:
                continue
            key = (loc.table_index, loc.row_index)
            groups.setdefault(key, {})[loc.col_index] = r.role
    return groups


def _set_cell_value(cell, value: str):
    """Zapiše vrednost v prvi run prve neprazne odstavka celice (ohrani obliko
    tega run-a - pisavo, velikost ...), ostale run-e v istem odstavku počisti."""
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = value
            try:
                p.runs[0].font.color.rgb = BLACK
            except Exception:
                pass
            for extra in p.runs[1:]:
                extra.text = ""
            return
    if cell.paragraphs:
        run = cell.paragraphs[0].add_run(value)
        run.font.color.rgb = BLACK


def expand_parcel_rows(doc, row_groups: Dict[Tuple[int, int], Dict[int, str]], parcels: List[dict]) -> int:
    """parcels: seznam slovarjev s ključi ime_priimek/naslov/katastrska_obcina/
    parcelna_stevilka/delez - parcels[0] je "predloga" (že zapisana v izvirni
    vrstici), za parcels[1:] ta funkcija DODA nove vrstice.
    Vrne število dejansko dodanih vrstic (na tabelo, seštej za skupno)."""
    added = 0
    if len(parcels) <= 1:
        return added

    for (t_idx, template_row_idx), col_roles in row_groups.items():
        if t_idx >= len(doc.tables):
            continue
        table = doc.tables[t_idx]
        if template_row_idx >= len(table.rows):
            continue
        template_tr = table.rows[template_row_idx]._tr
        anchor_tr = template_tr
        for parcel in parcels[1:]:
            new_tr = copy.deepcopy(template_tr)
            anchor_tr.addnext(new_tr)
            anchor_tr = new_tr
            new_row = _Row(new_tr, table)
            for col_idx, role in col_roles.items():
                if col_idx < len(new_row.cells):
                    _set_cell_value(new_row.cells[col_idx], parcel.get(role, ""))
            added += 1

    return added


def _clear_same_cell_duplicates(f: Field) -> int:
    """Če ima POLJE (f) več fizičnih pojavitev ZNOTRAJ ISTE celice tabele
    (isti table_index+row_index+col_index - torej gre za več ločenih
    odstavkov v ENI celici, ne za ločene vrstice), pusti vrednost samo v
    PRVI taki pojavitvi, preostale pa počisti (prazno besedilo) - da se ista
    vsebina ne izpiše dvakrat, druga pod drugo, v isti celici. Vrne število
    počiščenih odvečnih pojavitev."""
    by_cell: Dict[Tuple[int, int, int], List] = {}
    for loc in f.locations:
        if loc.table_index is None or loc.row_index is None:
            continue
        key = (loc.table_index, loc.row_index, loc.col_index or 0)
        by_cell.setdefault(key, []).append(loc)
    cleared = 0
    for locs in by_cell.values():
        for loc in locs[1:]:
            _set_location_value(loc, "")
            cleared += 1
    return cleared


def collapse_duplicate_parcel_cell_paragraphs(fields: Dict[str, Field], match_report, unique_parcels: List[dict]) -> int:
    """Če predloga za polja parcelna_stevilka/katastrska_obcina/delez v ISTI
    celici tabele vnaprej pripravi VEČ odstavkov (npr. zasnovano za primer
    do 2 RAZLIČNIH parcelnih številk), a je parcela med vsemi lastniki v
    resnici ENA SAMA (unique_parcels ima samo 1 vnos - npr. solastnika
    imata isto parcelo, vsak s svojim deležem), se ista vrednost sicer
    zapiše v VSE te vnaprej pripravljene odstavke iste celice in se zato
    podvojeno izpiše druga pod drugo (npr. "951/8" dvakrat namesto enkrat).
    Ta funkcija take odvečne podvojene odstavke znotraj iste celice počisti
    in pusti samo EN zapis.

    Deluje samo, kadar je unique_parcels dolg 1 - če je parcel dejansko VEČ
    (dolžina > 1), za to poskrbi `merge_parcels_single_row` (ki v isto
    celico zapiše seznam VSEH parcel, ločen z vejico); takrat te funkcije ne
    kličemo, da po nepotrebnem ne bi brisali legitimnih dodatnih vrstic.

    POMEMBNO: grupiranje po fizični celici (`_group_parcel_locations_by_cell`)
    zajame tudi primer, ko so v ISTI celici vnaprej pripravljeni odstavki z
    RAZLIČNIM vzorčnim besedilom (npr. '949/22' in '949/23' kot dve ločeni
    Field polji) - v tem primeru match_and_apply obe polji neodvisno
    zamenja z isto (edino) pravo vrednostjo, kar brez tega popravka pusti
    isto parcelno številko podvojeno (vsako v svojem odstavku iste celice).

    Vrne število dejansko počiščenih odvečnih pojavitev (za informativno
    sporočilo v UI)."""
    if len(unique_parcels) > 1:
        return 0
    cleared = 0
    by_cell = _group_parcel_locations_by_cell(fields, match_report)
    for cell_entries in by_cell.values():
        if len(cell_entries) <= 1:
            continue
        for _role, _f, loc in cell_entries[1:]:
            _set_location_value(loc, "")
            cleared += 1
    return cleared


def _set_location_value(loc, value: str):
    """Zapiše vrednost v runs ENE KONKRETNE lokacije polja (ne vseh pojavitev
    polja hkrati) in jo obarva črno - uporablja se za polja, ki se s to
    natančno besedilno vsebino pojavijo tako v prostem besedilu KOT v tabeli
    (docx_engine jih zaradi enakega besedila združi v isto polje), da lahko
    obe pojavitvi obravnavamo LOČENO."""
    if not loc.runs:
        return
    loc.runs[0].text = value
    try:
        loc.runs[0].font.color.rgb = BLACK
    except Exception:
        pass
    for extra_run in loc.runs[1:]:
        extra_run.text = ""
        try:
            extra_run.font.color.rgb = BLACK
        except Exception:
            pass


def expand_free_text_parcel_mentions(fields: Dict[str, Field], match_report, parcels: List[dict]) -> int:
    """Ko ima lastnik VEČ KOT ENO parcelo, poišče omembe parcelne številke
    (in po potrebi katastrske občine, če se ta med parcelami razlikuje) v
    PROSTEM BESEDILU predloge (izven tabel - npr. stavek "...da se na
    zemljišču s parc. št.: X, k.o. Y izrecno in nepogojno dovoljuje...") in
    jih nadomesti s SEZNAMOM VSEH parcel lastnika, kot so navedene v
    tabelah. Polja, ki se pojavijo IZKLJUČNO v tabelah, tu niso obravnavana -
    zanje poskrbi `expand_parcel_rows` (dodajanje novih vrstic).

    POMEMBNO: če je polje z enakim besedilom (npr. katastrska občina) hkrati
    tudi v eni ali več tabelah (ker gre za dobesedno isto besedilo), se
    posodobi SAMO lokacija v prostem besedilu - tabelske lokacije ostanejo
    nedotaknjene (te dobijo svoje vrednosti prek `expand_parcel_rows`).

    Vrne število dejansko posodobljenih omemb (za informativno sporočilo v UI)."""
    if len(parcels) <= 1:
        return 0

    all_numbers: List[str] = []
    seen_numbers = set()
    for p in parcels:
        num = (p.get("parcelna_stevilka") or "").strip()
        if num and num not in seen_numbers:
            seen_numbers.add(num)
            all_numbers.append(num)

    all_ko: List[str] = []
    seen_ko = set()
    for p in parcels:
        ko = (p.get("katastrska_obcina") or "").strip()
        if ko and ko not in seen_ko:
            seen_ko.add(ko)
            all_ko.append(ko)

    updated = 0
    for r in match_report:
        if r.role not in ("parcelna_stevilka", "katastrska_obcina"):
            continue
        f = fields.get(r.field_id)
        if f is None:
            continue

        free_text_locs = [loc for loc in f.locations if loc.table_index is None]
        if not free_text_locs:
            continue  # polje je samo v tabeli - to ureja expand_parcel_rows

        if r.role == "parcelna_stevilka" and len(all_numbers) > 1:
            # Ohrani morebitno končno ločilo iz predloge (npr. ", " pred
            # nadaljevanjem stavka "... k.o. ...") - poiščemo ga v obstoječi
            # (že zamenjani) vrednosti polja, ki vsebuje prvo/primarno parcelo.
            frac_m = _FRACTION_RE.search(f.value)
            suffix = f.value[frac_m.end():] if frac_m else ""
            new_text = ", ".join(all_numbers) + suffix
        elif r.role == "katastrska_obcina" and len(all_ko) > 1:
            new_text = ", ".join(all_ko)
        else:
            continue

        for loc in free_text_locs:
            _set_location_value(loc, new_text)
        updated += 1

    return updated


def dedupe_parcels(records: List[dict]) -> List[dict]:
    """Iz seznama zapisov (lahko vsebuje več vrstic za ISTO parcelo, npr. če
    ima parcela več solastnikov) izlušči seznam UNIKATNIH parcel, razvrščenih
    po vrstnem redu prvega pojava, prepoznanih po paru (katastrska_obcina,
    parcelna_stevilka)."""
    seen = set()
    parcels: List[dict] = []
    for rec in records:
        key = (
            (rec.get("katastrska_obcina") or "").strip().lower(),
            (rec.get("parcelna_stevilka") or "").strip().lower(),
        )
        if key not in seen:
            seen.add(key)
            parcels.append(rec)
    return parcels


def merge_parcels_single_row(fields: Dict[str, Field], match_report, parcels: List[dict]) -> int:
    """Namesto podvajanja vrstic za vsako dodatno parcelo istega lastnika
    (glej `expand_parcel_rows`), v OBSTOJEČO (edino) pojavitev polja zapiše
    SEZNAM VSEH parcelnih številk, ločen z vejico (npr. "949/22, 949/23") -
    uporabno pri predlogah, ki za več parcel istega lastnika predvidevajo ENO
    SKUPNO vrstico/celico namesto ločene vrstice za vsako parcelo (glej npr.
    stolpec "parc. št." v tabeli parcel).

    Ker `apply_field_value` novo vrednost zapiše v VSE pojavitve polja, se
    ista združena vrednost samodejno pojavi tudi v morebitnem prostem
    besedilu z enakim izvirnim besedilom (npr. "...s parc. št.: ...") - ni
    torej potrebe po ločeni obdelavi tabel in prostega besedila, kot pri
    `expand_parcel_rows` + `expand_free_text_parcel_mentions`.

    Katastrsko občino združi na enak način, če se med parcelami razlikuje -
    če je ista za vse parcele, ostane nespremenjena (že pravilno zapisana).
    Delež ostane nedotaknjen (privzeto ostane vrednost prve/primarne
    parcele) - v predlogah je za istega lastnika običajno naveden en skupni
    delež, ne po parceli posebej.

    POMEMBNO: če je predloga v ISTI celici vnaprej pripravila VEČ ločenih
    odstavkov z RAZLIČNIM vzorčnim besedilom (npr. '949/22' v prvem in
    '949/23' v drugem odstavku - torej DVE ločeni Field polji, ne eno z
    dvema lokacijama), ju `docx_engine.parse_docx_fields` ne združi v isto
    polje. Prejšnja različica te funkcije je klicala `apply_field_value` po
    enem polju naenkrat, kar take medsebojne podvojitve v isti celici ni
    opazilo - obe polji sta se neodvisno napolnili z isto vrednostjo in se
    ta izpisala podvojeno (dvakrat druga pod drugo). Zato tu namesto tega
    uporabimo `_group_parcel_locations_by_cell`, ki zajame VSE fizične
    pojavitve (ne glede na to, kateremu Field polju pripadajo) v isti
    celici - združen seznam zapišemo v PRVO, preostale pa počistimo.

    Vrne število dejansko posodobljenih FIZIČNIH CELIC (za informativno
    sporočilo v UI)."""
    if len(parcels) <= 1:
        return 0

    all_numbers: List[str] = []
    seen_numbers = set()
    for p in parcels:
        num = (p.get("parcelna_stevilka") or "").strip()
        if num and num not in seen_numbers:
            seen_numbers.add(num)
            all_numbers.append(num)

    all_ko: List[str] = []
    seen_ko = set()
    for p in parcels:
        ko = (p.get("katastrska_obcina") or "").strip()
        if ko and ko not in seen_ko:
            seen_ko.add(ko)
            all_ko.append(ko)

    by_cell = _group_parcel_locations_by_cell(fields, match_report)

    updated = 0
    for (_t_idx, _r_idx, _c_idx, role), cell_entries in by_cell.items():
        if role == "parcelna_stevilka" and len(all_numbers) > 1:
            new_text = ", ".join(all_numbers)
        elif role == "katastrska_obcina" and len(all_ko) > 1:
            new_text = ", ".join(all_ko)
        else:
            continue

        _role0, first_field, first_loc = cell_entries[0]
        _set_location_value(first_loc, new_text)
        first_field.value = new_text
        # Morebitne DRUGE prostobesedilne lokacije PRVEGA polja (izven
        # tabel, z enakim izvirnim besedilom kot ta celica) posodobimo
        # enako - glej opombo o samodejnem širjenju v prosto besedilo v
        # opisu funkcije zgoraj.
        for other_loc in first_field.locations:
            if other_loc is not first_loc and other_loc.table_index is None:
                _set_location_value(other_loc, new_text)

        # Preostale (podvojene) fizične pojavitve v ISTI celici - bodisi
        # dodatni odstavki istega polja, bodisi ločena Field polja z
        # drugačnim izvirnim vzorčnim besedilom - počistimo, da se
        # vrednost ne izpiše podvojeno.
        for _role_i, other_field, other_loc in cell_entries[1:]:
            _set_location_value(other_loc, "")
            if other_field is not first_field:
                other_field.value = ""

        updated += 1

    return updated
