"""
Nadzorna plošča za avtomatizacijo pravnih dokumentov: prebere PDF izpisek iz
zemljiške knjige (eZK), izlušči 5 ključnih podatkov (ime in priimek, naslov,
katastrska občina, parcelna številka, delež) in jih zamenja v ustreznih
rdečih poljih .docx predloge. Prikaže vizualno primerjavo pred/po.

Zagon:
    streamlit run app.py

Pred zagonom:
    pip install -r requirements.txt
    pip install requests beautifulsoup4

Nastavitev API ključa (da ga ni treba vsakič ročno vnašati):
    Ustvarite datoteko .env (kopijo .env.example) v isti mapi z vsebino:
        NVIDIA_API_KEY=vaš-ključ-tukaj
    Ključ ustvarite na https://build.nvidia.com (odprite poljuben model in kliknite "Get API Key").
    Datoteka .env NI del kode, ki jo delite naprej - drži jo samo lokalno.
"""

import base64
import json
import os
import re
import requests
from bs4 import BeautifulSoup
from pathlib import Path

import streamlit as st
import pandas as pd
import streamlit.components.v1 as components

try:
    from dotenv import load_dotenv, set_key
    load_dotenv()  # prebere .env, če obstaja (tiho ne naredi nič, če ne obstaja)
    DOTENV_AVAILABLE = True
except ImportError:
    DOTENV_AVAILABLE = False

from docx_engine import (
    DEFAULT_RED_HEX,
    apply_field_value,
    load_docx,
    parse_docx_fields,
    save_docx,
)
from pdf_engine import extract_pdf_text
from llm_engine import (
    LLMExtractionError,
    extract_land_registry_data,
    count_probable_owners,
    test_network_connectivity,
    fetch_available_models,
    DEFAULT_MODEL,
    DEFAULT_NVIDIA_BASE_URL,
    DEFAULT_GEMINI_BASE_URL,
    PROVIDER_NVIDIA,
    PROVIDER_GEMINI,
    PROVIDER_LABELS,
    PROVIDER_ENV_KEY_VARS,
    PROVIDER_ENV_BASE_URL_VARS,
    PROVIDER_FALLBACK_MODELS,
    detect_provider,
    base_url_for_provider,
)

@st.cache_data(show_spinner=False, ttl=86400)
def pridobi_podatke_obcine(ime_obcine: str):
    def normalize_name(name: str) -> str:
        return name.lower().replace("mestna občina ", "").replace("občina ", "").strip()
    
    url_seznam = "https://kdo-vodi.si/obcine"
    headers = {"User-Agent": "Mozilla/5.0"}
    
    try:
        res = requests.get(url_seznam, headers=headers, timeout=10)
        res.raise_for_status()
    except Exception:
        return None, None

    soup = BeautifulSoup(res.text, 'html.parser')
    table = soup.find('table')
    if not table:
        return None, None
        
    iskalno_ime = normalize_name(ime_obcine)
    zupan = None
    povezava_obcine = None
    
    for row in table.find_all('tr')[1:]:
        cols = row.find_all('td')
        if len(cols) >= 2:
            obcina_link_tag = cols[0].find('a')
            obcina_tekst = cols[0].get_text(strip=True)
            
            if normalize_name(obcina_tekst) == iskalno_ime:
                zupan = cols[1].get_text(strip=True)
                if obcina_link_tag and 'href' in obcina_link_tag.attrs:
                    povezava_obcine = obcina_link_tag['href']
                break

    if not zupan:
        return None, None

    maticna = "[MATIČNA ŠTEVILKA NI NAJDENA]"
    
    if povezava_obcine:
        if not povezava_obcine.startswith("http"):
            base_url = "https://kdo-vodi.si"
            povezava_obcine = base_url + povezava_obcine if povezava_obcine.startswith("/") else base_url + "/" + povezava_obcine
            
        try:
            res_pod = requests.get(povezava_obcine, headers=headers, timeout=10)
            res_pod.raise_for_status()
            soup_pod = BeautifulSoup(res_pod.text, 'html.parser')
            vidno_besedilo = soup_pod.get_text(separator=' ')
            match = re.search(r'MATIČNA ŠT\.?\s*(\d+)', vidno_besedilo, re.IGNORECASE)
            if match:
                maticna = match.group(1)
        except Exception:
            pass

    return zupan, maticna

@st.cache_data(show_spinner=False, ttl=3600)
def cached_extract_land_registry_data(pdf_text: str, api_key: str | None, proxy_url: str | None, model: str, base_url: str | None):
    return extract_land_registry_data(
        pdf_text=pdf_text,
        api_key=api_key,
        proxy_url=proxy_url,
        model=model,
        base_url=base_url
    )

@st.cache_data(show_spinner=False, ttl=3600)
def cached_fetch_ezk_pdf(_state: dict, katastrska_obcina: str, parcelna_stevilka: str, headless: bool) -> bytes:
    from ezk_engine import fetch_redni_izpis_pdf_with_state
    return fetch_redni_izpis_pdf_with_state(
        state=_state,
        katastrska_obcina=katastrska_obcina,
        parcelna_stevilka=parcelna_stevilka,
        headless=headless
    )


from matching_engine import match_and_apply, LAND_REGISTRY_KEYS
from multi_parcel_engine import dedupe_parcels, merge_parcels_single_row, collapse_duplicate_parcel_cell_paragraphs
from owner_engine import (
    dedupe_owners,
    assign_owners_to_slots,
    expand_owner_rows,
    expand_free_text_owner_mentions,
    fill_owner_list_summary_fields,
    )
import preview_engine
from owner_grouping_engine import (
    group_by_ownership_unit,
    group_by_katastrska_obcina,
    find_name_conflicts,
)

try:
    from ezk_engine import interactive_login, fetch_redni_izpis_pdf_with_state, EzkError, SEARCH_URL as EZK_SEARCH_URL
    EZK_ENGINE_AVAILABLE = True
except ImportError:
    EZK_ENGINE_AVAILABLE = False
    EZK_SEARCH_URL = "https://esodisce.si/evlozisce/javni_izpisi/list.html#"

from ko_registry import get_ko_name

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

# --------------------------------------------------------------------------
# SAMODEJNA IZBIRA PREDLOGE POGODBE (mapa templates/ ob app.py)
# --------------------------------------------------------------------------
TEMPLATES_DIR = Path(__file__).parent / "templates"

TEMPLATE_1_MOSKI = ("1_lastnik_moski.docx", "1 lastnik (moški)")
TEMPLATE_1_ZENSKA = ("1_lastnik_zenska.docx", "1 lastnica (ženska)")
TEMPLATE_2 = ("2_lastnika.docx", "2 lastnika")
TEMPLATE_3 = ("3_lastniki.docx", "3 lastniki")
TEMPLATE_4 = ("4_lastniki.docx", "4 lastniki")
TEMPLATE_5PLUS = ("5_lastnikov.docx", "5 ali več lastnikov")
TEMPLATE_OBCINA = ("obcina.docx", "Občina (lastnik je občina)")

_MALE_NAME_EXCEPTIONS = {"luka", "jaka", "matija", "ilija", "nikola", "andrija"}


def _first_name(ime_priimek: str) -> str:
    parts = (ime_priimek or "").strip().split()
    return parts[0] if parts else ""


def guess_owner_gender(owner: dict) -> str | None:
    for key in ("spol", "gender", "sex"):
        val = (owner.get(key) or "").strip().lower()
        if val in ("m", "moški", "moski", "male"):
            return "M"
        if val in ("z", "ž", "ženska", "zenska", "female"):
            return "Z"

    first = _first_name(owner.get("ime_priimek", "")).lower()
    if not first:
        return None
    if first in _MALE_NAME_EXCEPTIONS:
        return "M"
    if first.endswith("a"):
        return "Z"
    return "M"


def pick_contract_template(unique_owners: list[dict]) -> tuple[str, str, bool]:
    n = len(unique_owners)
    if n == 1:
        owner = unique_owners[0]
        name = (owner.get("ime_priimek") or "").lower()
        
        if "občina" in name or "obcina" in name:
            fname, label = TEMPLATE_OBCINA
            return fname, label, False

        gender = guess_owner_gender(owner)
        has_explicit = any((owner.get(k) or "").strip() for k in ("spol", "gender", "sex"))
        uncertain = not has_explicit
        if gender == "Z":
            fname, label = TEMPLATE_1_ZENSKA
        else:
            fname, label = TEMPLATE_1_MOSKI
        return fname, label, uncertain
    elif n == 2:
        fname, label = TEMPLATE_2
        return fname, label, False
    elif n == 3:
        fname, label = TEMPLATE_3
        return fname, label, False
    elif n == 4:
        fname, label = TEMPLATE_4
        return fname, label, False
    else:
        fname, label = TEMPLATE_5PLUS
        return fname, label, False


@st.cache_data(show_spinner=False)
def _load_template_bytes(filename: str) -> bytes | None:
    try:
        return (TEMPLATES_DIR / filename).read_bytes()
    except Exception:
        return None


def _activate_template(file_bytes: bytes, file_id: str) -> None:
    st.session_state.original_docx_bytes = file_bytes
    doc = load_docx(file_bytes)
    st.session_state.doc = doc
    st.session_state.fields = parse_docx_fields(doc, red_hex=st.session_state.red_hex_codes)
    st.session_state.template_file_id = file_id
    st.session_state.match_report = None
    st.session_state.field_roles = {}
    st.session_state.field_changes = {}
    st.session_state.print_pdf_bytes = None


# --------------------------------------------------------------------------
# LOGOTIP PODJETJA (zgornji desni kot)
# --------------------------------------------------------------------------
LOGO_PATH = Path(__file__).parent / "assets" / ""

def _load_logo_b64() -> str | None:
    try:
        return base64.b64encode(LOGO_PATH.read_bytes()).decode("utf-8")
    except Exception:
        return None

_LOGO_B64 = _load_logo_b64()

# --------------------------------------------------------------------------
# LOGOTIP SI-PASS
# --------------------------------------------------------------------------
SIPASS_LOGO_PATH = Path(__file__).parent / "assets" / "sipass.png"

def _load_sipass_logo_b64() -> str | None:
    try:
        return base64.b64encode(SIPASS_LOGO_PATH.read_bytes()).decode("utf-8")
    except Exception:
        return None

_SIPASS_LOGO_B64 = _load_sipass_logo_b64()

# --------------------------------------------------------------------------
# PRIVZET MODEL/NASLOV (API naslov + model)
# --------------------------------------------------------------------------
DEFAULT_SETTINGS_PATH = Path(__file__).parent / "default_api_settings.json"

def _load_default_settings() -> dict:
    try:
        return json.loads(DEFAULT_SETTINGS_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _write_default_settings(base_url: str | None = None, model: str | None = None, provider: str | None = None):
    try:
        data = {}
        if DEFAULT_SETTINGS_PATH.exists():
            try:
                data = json.loads(DEFAULT_SETTINGS_PATH.read_text(encoding="utf-8"))
            except Exception:
                pass
        
        if base_url is not None:
            data["base_url"] = base_url
        if model is not None:
            data["model"] = model
        if provider is not None:
            data["provider"] = provider

        DEFAULT_SETTINGS_PATH.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        pass


# --------------------------------------------------------------------------
# SHRANJEVANJE API KLJUČA + NASLOVA V .env
# --------------------------------------------------------------------------
ENV_FILE_PATH = Path(__file__).parent / ".env"

def _save_provider_credentials_to_env(provider: str, api_key: str, base_url: str) -> None:
    if not DOTENV_AVAILABLE:
        raise RuntimeError("knjižnica 'python-dotenv' ni nameščena (pip install python-dotenv)")
    if not api_key:
        raise RuntimeError("API ključ je prazen - ni česa shraniti")

    if not ENV_FILE_PATH.exists():
        ENV_FILE_PATH.touch()

    key_var = PROVIDER_ENV_KEY_VARS.get(provider, "NVIDIA_API_KEY")
    base_url_var = PROVIDER_ENV_BASE_URL_VARS.get(provider, "NVIDIA_BASE_URL")

    set_key(str(ENV_FILE_PATH), key_var, api_key)
    if base_url:
        set_key(str(ENV_FILE_PATH), base_url_var, base_url)


# --------------------------------------------------------------------------
# OSNOVNE NASTAVITVE STRANI + CSS
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="e-Služnost",
    page_icon="assets/logo.png",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    :root {
        --primary-blue: #1E4A7B;
        --bg-gray: #F5F7FA;
        --border-gray: #E2E6EB;
        --text-dark: #1F2937;
        --ok-green: #1E8E3E;
        --warn-amber: #B98600;
    }
    .main { background-color: var(--bg-gray); }
    h1, h2, h3 { color: var(--primary-blue); font-weight: 600; }
    .block-container { padding-top: 1.5rem; }
    .status-pill {
        display: inline-block; padding: 3px 10px; border-radius: 999px;
        font-size: 0.8rem; font-weight: 600; margin-bottom: 4px; margin-right: 4px;
    }
    .pill-ok { background-color: #E4F4E8; color: var(--ok-green); }
    .pill-warn { background-color: #FBF0DC; color: var(--warn-amber); }
    .pill-neutral { background-color: #E9EDF3; color: var(--primary-blue); }
    .pill-changed { background-color: #FDECEC; color: #B3261E; }
    .field-caption { color: #6B7280; font-size: 0.78rem; margin-top: -8px; }
    .stTabs [data-baseweb="tab-list"] {
        gap: 12px;
        display: flex;
        align-items: center;
        justify-content: flex-start;
        flex-wrap: nowrap;
        width: 100%;
        padding: 2px 0 18px 0;
        margin-bottom: 2px;
        border-bottom: 1px solid rgba(30, 74, 123, 0.12);
    }
    .stTabs [data-baseweb="tab"] {
        height: auto;
        min-height: 54px;
        padding: 12px 26px;
        border-radius: 14px;
        border: 1px solid #B8C7D9;
        background: linear-gradient(180deg, #FFFFFF 0%, #EEF4F9 100%);
        box-shadow: inset 0 0 0 1px rgba(255,255,255,0.85), 0 4px 10px rgba(29, 78, 116, 0.08);
        font-weight: 900;
        color: var(--primary-blue);
        margin-right: 6px;
        transition: all 180ms ease;
    }
    .stTabs [data-baseweb="tab"] p {
        font-size: 1.14rem !important;
        line-height: 1.24;
        font-weight: 900 !important;
        letter-spacing: 0.05em;
        text-transform: uppercase;
        color: var(--primary-blue);
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        border-color: var(--primary-blue);
        background: linear-gradient(180deg, #1E4A7B 0%, #244F7C 100%);
        box-shadow: 0 4px 14px rgba(30, 74, 123, 0.30), inset 0 0 0 1px rgba(255,255,255,0.38);
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] p {
        color: white !important;
    }
    .stTabs [data-baseweb="tab"]:hover {
        border-color: var(--primary-blue);
        background: linear-gradient(180deg, #FFFFFF 0%, #DDECF6 100%);
        transform: translateY(-2px);
    }
    .stTabs [data-baseweb="tab"]:nth-child(3) {
        margin-left: auto;
    }
    .stTabs [data-baseweb="tab"]:nth-child(2) {
        font-size: 1.34rem !important;
        padding-left: 30px;
        padding-right: 30px;
    }
    .stTabs [data-baseweb="tab"]:first-child {
        min-width: unset;
    }
    .em-logo {
        position: fixed !important;
        top: 40px;
        right: 32px;
        z-index: 999999;
        width: 240px;
        height: auto;
        pointer-events: none;
    }
    @media (max-width: 900px) {
        .em-logo { width: 150px; top: 42px; right: 14px; }
    }
    [data-testid="stAppViewContainer"],
    [data-testid="stMain"],
    [data-testid="stMainBlockContainer"],
    [data-testid="stBottomBlockContainer"],
    .main,
    .block-container,
    .stApp {
        transform: none !important;
        filter: none !important;
        perspective: none !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
    
)

if _LOGO_B64:
    st.markdown(
        f'<img id="logo-img" src="data:image/jpeg;base64,{_LOGO_B64}" '
        f'class="logo" alt="Elektro Maribor d.d.">',
        unsafe_allow_html=True,
    )
    components.html(
        """
        <script>
        try {
            var doc = window.parent.document;
            var img = doc.getElementById('em-logo-img');
            if (img && img.parentElement !== doc.body) {
                doc.body.appendChild(img);
            }
        } catch (e) { /* tiho prezri - CSS popravek zgoraj naj zadostuje */ }
        </script>
        """,
        height=0,
        width=0,
    )

st.title("")
st.image("assets/naslovna.png", width=3000)

st.caption(
    "Aplikacija je namenjena izključno za interno uporabo v podjetju Elektro Maribor d.d. in je zaščitena z avtorskimi pravicami. Vsebine, pridobljene z uporabo aplikacije, so zaupne in se ne smejo deliti ali objavljati zunaj podjetja brez ustreznega dovoljenja."
)

# --------------------------------------------------------------------------
# SEJNO STANJE ZA PRIJAVO (SI-PASS)
# --------------------------------------------------------------------------
ezk_defaults = {
    "ezk_logged_in": False,
    "ezk_username": "",
    "ezk_storage_state": None,
    "ezk_last_pdf": None,
    "ezk_debug_mode": False,
}
for k, v in ezk_defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

if not st.session_state.ezk_logged_in:
    st.markdown(
        "<div style='max-width: 460px; margin: 50px auto 0 auto; text-align: left;'>",
        unsafe_allow_html=True,
    )
    st.markdown(
        "<p style='font-weight: 700; text-transform: uppercase; font-size: 1.05rem; "
        "letter-spacing: 0.02em; margin-bottom: 4px;'>"
        "PRIJAVA POTEKA IZKLJUČNO PREKO SI-PASS S KVALIFICIRANIM DIGITALNIM POTRDILOM - "
        "KLIKNITE NA LOGOTIP SPODAJ.</p>",
        unsafe_allow_html=True,
    )

    if not EZK_ENGINE_AVAILABLE:
        st.error(
            "Manjka knjižnica 'playwright', zato prijava ne deluje. Namestite jo z:\n\n"
            "`pip install playwright --break-system-packages`\n\n`playwright install chromium`"
        )
    else:
        _sipass_primary_src = (
            f"data:image/png;base64,{_SIPASS_LOGO_B64}"
            if _SIPASS_LOGO_B64
            else "https://esodisce.si/spa/assets/images/si-pass-logo.png"
        )
        st.markdown(
            f'''
            <div id="sipass-login-trigger" style="display: inline-block; line-height: 0; margin: 15px 0 25px 0; cursor: pointer;"
                 title="Klikni za prijavo preko SI-PASS">
                <img src="{_sipass_primary_src}"
                     alt="SI-PASS - klikni za prijavo" style="height: 70px; max-width: 100%; border-radius: 8px;"
                     onerror="this.onerror=null; this.src='https://www.si-trust.gov.si/assets/Uploads/Logotipi/SI-PASS-logo-v2.png';">
            </div>
            <style>
            div:has(> #sipass-login-btn-anchor) + div {{ display: none !important; }}
            </style>
            <div id="sipass-login-btn-anchor"></div>
            ''', unsafe_allow_html=True
        )
        sipass_login_clicked = st.button(
            "🔑 Prijava preko SI-PASS (Kvalificirano potrdilo)", use_container_width=True, type="primary",
            key="sipass_login_btn",
        )
        components.html(
            """
            <script>
            (function() {
                var attempts = 0;
                function wire() {
                    attempts++;
                    try {
                        var doc = window.parent.document;
                        var trigger = doc.getElementById('sipass-login-trigger');
                        var buttons = doc.querySelectorAll('button');
                        var loginBtn = null;
                        for (var i = 0; i < buttons.length; i++) {
                            var txt = buttons[i].innerText || buttons[i].textContent || "";
                            if (txt.indexOf("Prijava preko SI-PASS") !== -1) {
                                loginBtn = buttons[i];
                                break;
                            }
                        }
                        if (trigger && loginBtn && !trigger.dataset.wired) {
                            trigger.dataset.wired = "1";
                            trigger.addEventListener('click', function() {
                                loginBtn.click();
                            });
                            var btnWrapper = loginBtn.closest('[data-testid="stButton"]') || loginBtn.parentElement;
                            if (btnWrapper) { btnWrapper.style.display = 'none'; }
                            return;
                        }
                    } catch (e) { /* tiho prezri */ }
                    if (attempts < 30) {
                        setTimeout(wire, 150);
                    }
                }
                wire();
            })();
            </script>
            """,
            height=0, width=0,
        )

        if sipass_login_clicked:
            try:
                with st.spinner("Odpiram brskalnik za prijavo... Prosim, potrdite certifikat v novem oknu."):
                    from ezk_engine import interactive_login
                    username, state = interactive_login()
                    st.session_state.ezk_username = username
                    st.session_state.ezk_storage_state = state
                    st.session_state.ezk_logged_in = True
                    
                    # ---- NOVO: Ohranjanje seje pri življenju v ozadju ----
                    def keep_alive_worker(state, max_minutes=30, interval_minutes=9):
                        import time
                        import asyncio
                        import threading
                        from playwright.sync_api import sync_playwright
                        
                        # Da Playwright deluje v ločeni niti, potrebujemo nov event loop
                        asyncio.set_event_loop(asyncio.new_event_loop())
                        start_time = time.time()
                        
                        while (time.time() - start_time) < (max_minutes * 60):
                            time.sleep(interval_minutes * 60)
                            try:
                                with sync_playwright() as p:
                                    browser = p.chromium.launch(headless=True)
                                    context = browser.new_context(storage_state=state)
                                    page = context.new_page()
                                    # Kratek ping na eSodstvo, da se seja osveži
                                    page.goto(EZK_SEARCH_URL, timeout=15000)
                                    browser.close()
                            except Exception:
                                pass # Tiho ignoriraj napake pri osveževanju v ozadju

                    import threading
                    t = threading.Thread(
                        target=keep_alive_worker, 
                        args=(state, 30, 9), 
                        daemon=True
                    )
                    t.start()
                    # -------------------------------------------------------

                    st.rerun()
            except EzkError as e:
                st.error(str(e))
            except Exception as e:
                print(f"[SI-PASS prijava] Nepričakovana napaka: {e}")
                st.error("NAPAKA PRI PRIJAVI! POSKUSI ZNOVA")

    st.markdown("</div>", unsafe_allow_html=True)
    st.stop()

# --------------------------------------------------------------------------
# SEJNO STANJE
# --------------------------------------------------------------------------
defaults = {
    "doc": None,
    "original_docx_bytes": None,
    "fields": {},
    "template_file_id": None,
    "auto_template_info": None,
    "pdf_records": [],
    "red_hex_codes": list(DEFAULT_RED_HEX),
    "match_report": None,
    "field_roles": {},
    "field_changes": {},
    "parcels_data": None,
    "print_pdf_bytes": None,
    "transferred_parcels": None,
    "transferred_owner_label": None,
    "transferred_zupan": None,
    "transferred_maticna": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

if "pdf_uploader_version" not in st.session_state:
    st.session_state.pdf_uploader_version = 0

# --------------------------------------------------------------------------
# SEJNO STANJE - ZAVIHEK "PROJEKT"
# --------------------------------------------------------------------------
if "projekt_info" not in st.session_state:
    st.session_state.projekt_info = {
        "naziv_gradnje": "",
        "st_dokumentacije": "",
        "datum_izdelave": "",
    }

def reset_all():
    for k, v in defaults.items():
        st.session_state[k] = v
    st.session_state.pdf_uploader_version += 1

def _current_field_text(f) -> str:
    if not f.locations:
        return f.value
    loc = f.locations[0]
    if not loc.runs:
        return f.value
    return "".join(r.text for r in loc.runs)

def sync_fields_from_widgets():
    for fid, f in st.session_state.fields.items():
        widget_key = f"field_input_{fid}"
        if widget_key in st.session_state:
            new_val = st.session_state[widget_key]
            if new_val != _current_field_text(f):
                apply_field_value(f, new_val)

def parse_ko_parcels_from_text(text: str):
    import re
    results = []
    
    # 1. VAROVALO: Popravimo pogoste OCR napake (če prebere črko O namesto številke 0 pri šifri KO)
    text = re.sub(r'\b[oO](\d{3,4})\b', r'0\1', text)

    # 1b. VAROVALO: OCR pogosto izpusti piko za zaporedno številko ("4. 0635" -> "40635"),
    # zaradi česar se zap. št. zlije s 4-mestno šifro KO v eno 5- ali 6-mestno število,
    # ki ga spodnji vzorci ne prepoznajo (šifra KO mora biti natanko 3-4 mestna).
    # Tu ju spet ločimo: zadnje 4 številke obravnavamo kot šifro KO, preostanek spredaj
    # (1-2 številki) kot zaporedno številko.
    text = re.sub(r'(?<!\d)(\d{1,2})(\d{4})(?!\d)', r'\1 \2', text)
    
    # Eksplicitni vzorec (za klasično besedilo v odstavkih)
    pattern_explicit = re.finditer(r'(?:ko|k\.o\.|sifra|šifra)\s*[:\-]?\s*(\d{3,4}).*?(?:parc|parcela|p\.š\.|st\.|št\.)\s*[:\-]?\s*(\d+(?:/\d+)?)', text, re.IGNORECASE | re.DOTALL)
    for m in pattern_explicit:
        results.append((m.group(1), m.group(2)))
    
    if not results:
        # Vzorec po vrsticah (za branje iz tabel)
        for line in text.split('\n'):
            # Če je OCR naredil presledek znotraj parcele (npr. "433 / 2"), ga združimo v "433/2"
            line = re.sub(r'(\d+)\s*/\s*(\d+)', r'\1/\2', line)
            
            # Odstranimo zaporedno številko s piko na začetku (npr "1.", "12 .")
            line_clean = re.sub(r'^\s*\d+\s*\.\s*', '', line)
            
            # Poiščemo vse številke (cele in tiste z ulomki)
            numbers = re.findall(r'(?<!\d)\d+(?:/\d+)?(?!\d)', line_clean)
            
            if len(numbers) >= 2:
                # Varovalo, če OCR ni prepoznal pike za zap.št. (npr. prebere "1 0635 433/2")
                if len(numbers) >= 3 and '/' not in numbers[0] and len(numbers[0]) <= 2 and len(numbers[1]) >= 3:
                    numbers = numbers[1:] # Odstranimo prvo številko, ker je očitno zaporedna
                    
                ko_candidate = None
                parc_candidate = None
                
                for i, num in enumerate(numbers):
                    # KO je običajno 3- ali 4-mestna številka
                    if re.fullmatch(r'\d{3,4}', num):
                        ko_candidate = num
                        
                        remaining = numbers[i+1:]
                        if remaining:
                            # Preferiramo številko s poševnico (značilno za parcele)
                            with_slash = [n for n in remaining if '/' in n]
                            if with_slash:
                                parc_candidate = with_slash[0]
                            else:
                                # Če poševnice ni (npr. parcela 438), vzamemo prvo naslednjo številko
                                parc_candidate = remaining[0]
                        break
                
                if ko_candidate and parc_candidate and ko_candidate != parc_candidate:
                    results.append((ko_candidate, parc_candidate))
                    
    # Deduplikacija (odstranimo dvojnike, a ohranimo vrstni red)
    seen = set()
    unique = []
    for ko, parc in results:
        if (ko, parc) not in seen:
            seen.add((ko, parc))
            unique.append((ko, parc))
    return unique

def _safe_filename_part(text: str) -> str:
    import re as _re
    import unicodedata as _ud
    text = _ud.normalize("NFKD", text or "").encode("ascii", "ignore").decode("ascii")
    text = _re.sub(r"\s+", "_", text.strip())
    text = _re.sub(r"[^A-Za-z0-9_\-]", "", text)
    return text

def build_export_filename() -> str:
    parcels = st.session_state.get("parcels_data")
    name_part = ""
    if parcels:
        primary = parcels[0]
        name_part = (
            primary.get("obcina")
            or primary.get("ime_obcine")
            or primary.get("ime_priimek")
            or ""
        )
    safe = _safe_filename_part(name_part)
    return f"pogodba_{safe}.docx" if safe else "pogodba_izpolnjena.docx"

LABELS_SL = {
    "ime_priimek": "Ime in priimek",
    "naslov": "Naslov",
    "katastrska_obcina": "Katastrska občina",
    "parcelna_stevilka": "Parcelna številka",
    "delez": "Delež",
    "ime_priimek+naslov": "Ime, priimek in naslov",
}

# --------------------------------------------------------------------------
# ZAVIHEK "PROJEKT" - samodejno prepisovanje v tabelo pogodbe
# (Naziv gradnje / Št. dokumentacije / Datum izdelave)
# --------------------------------------------------------------------------
PROJEKT_LABELS_SL = {
    "naziv_gradnje": "Naziv gradnje",
    "st_dokumentacije": "Št. dokumentacije",
    "datum_izdelave": "Datum izdelave",
}


def _normalize_label(text: str) -> str:
    import unicodedata as _ud
    text = _ud.normalize("NFKD", text or "")
    text = "".join(c for c in text if not _ud.combining(c))
    return text.lower().strip()


def _match_projekt_key(label_text: str) -> str | None:
    """Prepozna, ali oznaka v levem stolpcu tabele ustreza enemu od treh
    podatkov o projektu (naziv gradnje / št. dokumentacije / datum izdelave)."""
    norm = _normalize_label(label_text).rstrip(":").strip()
    if not norm:
        return None
    if "naziv gradnj" in norm:
        return "naziv_gradnje"
    if "dokumentacij" in norm:
        return "st_dokumentacije"
    if "datum" in norm and "izdelav" in norm:
        return "datum_izdelave"
    return None


def apply_projekt_fields_to_doc(doc, fields: dict, projekt_info: dict) -> int:
    """Poišče v tabelah dokumenta vrstice z oznakami 'Naziv gradnje:',
    'Št. dokumentacije:' in 'Datum izdelave:' (glej priloženo sliko) ter v
    isto vrstico (rdeče besedilo v drugem stolpcu) vpiše podatke iz zavihka
    'PROJEKT'. Ujemajoča polja odstrani iz `fields`, da jih kasnejši postopek
    ujemanja (match_and_apply) ne prepiše nazaj."""
    if not projekt_info:
        return 0

    applied = 0
    fields_to_remove = []

    for fid, f in fields.items():
        matched_key = None
        for loc in f.locations:
            if loc.table_index is None or loc.row_index is None:
                continue
            try:
                table = doc.tables[loc.table_index]
                row = table.rows[loc.row_index]
                label_text = row.cells[0].text if row.cells else ""
            except IndexError:
                continue
            key = _match_projekt_key(label_text)
            if key:
                matched_key = key
                break

        if matched_key and projekt_info.get(matched_key):
            apply_field_value(f, projekt_info[matched_key])
            applied += 1
            fields_to_remove.append(fid)

    for fid in fields_to_remove:
        fields.pop(fid, None)

    return applied

# --------------------------------------------------------------------------
# SEJNO STANJE - ZAVIHEK "SKUPINSKI PREGLED LASTNIKOV"
# --------------------------------------------------------------------------
MAX_BULK_PDFS = 30

bulk_defaults = {
    "bulk_pdf_records": [],
    "bulk_extracted": None,
    "bulk_owners": None,
    "bulk_by_ko": None,
}
for k, v in bulk_defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

if "bulk_uploader_version" not in st.session_state:
    st.session_state.bulk_uploader_version = 0

# --------------------------------------------------------------------------
# PRIVZETE VREDNOSTI ZA API NASTAVITVE (na voljo tudi, če je zavihek
# "Nastavitve" zaklenjen in ga uporabnik še ni odklenil)
# --------------------------------------------------------------------------
def _resolve_startup_api_defaults():
    """Ponovi enako logiko izbire ključa/naslova kot zavihek 'Nastavitve',
    da se API ključ vedno pravilno ujema s pripadajočim naslovom (base_url) -
    tudi preden administrator kadarkoli odklene nastavitve."""
    settings = _load_default_settings()
    preferred_provider = settings.get("provider")

    nvidia_key_env = (os.environ.get("NVIDIA_API_KEY") or "").strip()
    gemini_key_env = (os.environ.get("GEMINI_API_KEY") or "").strip()

    options = []
    if nvidia_key_env:
        options.append({
            "provider": PROVIDER_NVIDIA,
            "key": nvidia_key_env,
            "base_url": (os.environ.get("NVIDIA_BASE_URL") or "").strip() or DEFAULT_NVIDIA_BASE_URL,
        })
    if gemini_key_env:
        options.append({
            "provider": PROVIDER_GEMINI,
            "key": gemini_key_env,
            "base_url": (os.environ.get("GEMINI_BASE_URL") or "").strip() or DEFAULT_GEMINI_BASE_URL,
        })

    chosen = None
    if options:
        chosen = options[0]
        if preferred_provider:
            for opt in options:
                if opt["provider"] == preferred_provider:
                    chosen = opt
                    break

    api_key = chosen["key"] if chosen else ""
    desired_base_url = chosen["base_url"] if chosen else ""

    default_base_url = (settings.get("base_url") or "").strip()
    base_url = default_base_url or desired_base_url or DEFAULT_NVIDIA_BASE_URL
    model = (settings.get("model") or "").strip() or DEFAULT_MODEL

    return api_key, base_url, model


if "current_api_key" not in st.session_state:
    _startup_key, _startup_base_url, _startup_model = _resolve_startup_api_defaults()
    st.session_state.current_api_key = _startup_key
    st.session_state.current_proxy_url = os.environ.get("HTTPS_PROXY", os.environ.get("HTTP_PROXY", ""))
    st.session_state.current_base_url = _startup_base_url
    st.session_state.current_model = _startup_model

# --------------------------------------------------------------------------
# ZAVIHKI
# --------------------------------------------------------------------------
if "switch_to_fill_tab" not in st.session_state:
    st.session_state.switch_to_fill_tab = False

TAB_EZK_LABEL = "🏛️ E-ZEMLJIŠKA KNJIGA"
TAB_PROJEKT_LABEL = "🏗️ PROJEKT"
TAB_FILL_LABEL = "📝 IZPOLNJEVANJE POGODBE"
TAB_BULK_LABEL = f"📊 PREGLED LASTNIKOV PARCEL IZ ZEMLJIŠKE KNJIGE"
TAB_SETTINGS_LABEL = "⚙️ NASTAVITVE"

tab_ezk, tab_projekt, tab_bulk, tab_fill, tab_settings = st.tabs([
    TAB_EZK_LABEL,
    TAB_PROJEKT_LABEL,
    TAB_BULK_LABEL,
    TAB_FILL_LABEL,
    TAB_SETTINGS_LABEL,
])

if st.session_state.switch_to_fill_tab:
    st.session_state.switch_to_fill_tab = False
    components.html(
        """
        <script>
        (function() {
            var attempts = 0;
            function clickTab() {
                attempts++;
                try {
                    var doc = window.parent.document;
                    var tabs = doc.querySelectorAll('[data-baseweb="tab"]');
                    for (var i = 0; i < tabs.length; i++) {
                        var txt = tabs[i].innerText || tabs[i].textContent || "";
                        if (txt.indexOf("IZPOLNJEVANJE POGODBE") !== -1) {
                            tabs[i].click();
                            scrollFillTabToTop();
                            return;
                        }
                    }
                } catch (e) { /* tiho prezri */ }
                if (attempts < 30) {
                    setTimeout(clickTab, 100);
                }
            }
            function scrollFillTabToTop() {
                var tries = 0;
                function doScroll() {
                    tries++;
                    try {
                        var doc = window.parent.document;
                        window.parent.scrollTo({ top: 0, left: 0, behavior: "auto" });
                        if (doc.documentElement) { doc.documentElement.scrollTop = 0; }
                        if (doc.body) { doc.body.scrollTop = 0; }
                        var scrollTargets = doc.querySelectorAll(
                            '[data-testid="stMain"], [data-testid="stAppViewContainer"], section.main'
                        );
                        for (var j = 0; j < scrollTargets.length; j++) {
                            scrollTargets[j].scrollTop = 0;
                        }
                    } catch (e) { /* tiho prezri */ }
                    if (tries < 10) {
                        setTimeout(doScroll, 80);
                    }
                }
                doScroll();
            }
            clickTab();
        })();
        </script>
        """,
        height=0, width=0,
    )

with tab_ezk:
    st.subheader("🏛️ Iskalnik javnih izpisov iz zemljiške knjige (eSodstvo)")
    st.caption(
        f"Prenesen PDF se samodejno doda tudi v zavihek '{TAB_BULK_LABEL}'."
    )

    logout_col, _spacer_col = st.columns([1, 4])
    with logout_col:
        if st.button("🚪 Odjava (SI-PASS)", key="ezk_logout", help="Zapri sejo in odstrani piškotke.", use_container_width=True):
            st.session_state.ezk_logged_in = False
            st.session_state.ezk_username = ""
            st.session_state.ezk_storage_state = None
            st.rerun()

    st.markdown("**03-001 - Redni izpis iz zemljiške knjige**")
    with st.container(border=True):
        c1, c2 = st.columns(2)
        with c1:
            st.selectbox(
                "Način vnosa nepremičnin:", options=["po ID znaku"], disabled=True,
            )
        with c2:
            st.selectbox(
                "Tip nepremičnine:", options=["zemljiška parcela"], disabled=True,
            )
        st.markdown(
            '<div class="field-caption">✏️ Vpišite šifro katastrske občine in parcelno '
            'številko spodaj (ime KO se izpiše sproti).</div>',
            unsafe_allow_html=True,
        )

        k1, k2 = st.columns(2)
        with k1:
            ezk_katastrska_obcina = st.text_input("Katastrska občina (šifra)", key="ezk_katastrska_obcina")
            if ezk_katastrska_obcina:
                st.markdown(f"📍 **Mesto:** {get_ko_name(ezk_katastrska_obcina)}")

        with k2:
            ezk_parcelna_stevilka = st.text_input("Parcelna številka", key="ezk_parcelna_stevilka")

        st.session_state.ezk_debug_mode = st.checkbox(
            "🐞 Debug način (prikaži brskalnik med iskanjem)",
            value=st.session_state.ezk_debug_mode,
        )

        ezk_submit_disabled = not (
            ezk_katastrska_obcina.strip() and ezk_parcelna_stevilka.strip()
        )
        
        btn_col1, btn_col2 = st.columns(2)
        with btn_col1:
            btn_main = st.button(
                "📄 Prikaži v pdf obliki (dodaj v pregled)", type="primary", use_container_width=True,
                disabled=ezk_submit_disabled, key="ezk_submit"
            )
        with btn_col2:
            btn_direct = st.button(
                "📥 Samo prenesi PDF", type="secondary", use_container_width=True,
                disabled=ezk_submit_disabled, key="ezk_submit_direct"
            )

        st.link_button(
            "🔗 Odpri iskalnik eSodstvo neposredno",
            EZK_SEARCH_URL,
            use_container_width=True,
            help="Odpre uradni iskalnik javnih izpisov eSodstva v novem zavihku brskalnika - "
                 "za ročno iskanje (npr. po imenu katastrske občine), namesto samodejnega iskanja zgoraj.",
        )

        if btn_main or btn_direct:
            try:
                with st.spinner("Iskanje izpiska v eSodstvu ... (lahko traja nekaj sekund)"):
                    pdf_bytes = cached_fetch_ezk_pdf(
                        _state=st.session_state.ezk_storage_state,
                        katastrska_obcina=ezk_katastrska_obcina.strip(),
                        parcelna_stevilka=ezk_parcelna_stevilka.strip(),
                        headless=not st.session_state.ezk_debug_mode,
                    )
                file_name = f"ezk_{ezk_katastrska_obcina.strip()}_{ezk_parcelna_stevilka.strip()}.pdf"
                file_name = _safe_filename_part(file_name.replace(".pdf", "")) + ".pdf"
                st.session_state.ezk_last_pdf = (file_name, pdf_bytes)

                if btn_main:
                    file_id = f"{file_name}_{len(pdf_bytes)}"
                    existing_ids = {r["file_id"] for r in st.session_state.bulk_pdf_records}
                    if file_id not in existing_ids:
                        if len(st.session_state.bulk_pdf_records) < MAX_BULK_PDFS:
                            st.session_state.bulk_pdf_records.append(
                                {"file_id": file_id, "name": file_name, "text": extract_pdf_text(pdf_bytes)}
                            )
                        else:
                            st.warning(f"Zavihek '{TAB_BULK_LABEL}' je že poln.")
                    st.success(
                        f"Izpisek uspešno pridobljen in dodan v '{TAB_BULK_LABEL}' "
                        f"({len(pdf_bytes) / 1024:.0f} KB)."
                    )
                else:
                    import tkinter as tk
                    from tkinter import filedialog
                    
                    root = tk.Tk()
                    root.withdraw()
                    root.wm_attributes('-topmost', 1)
                    
                    saved_path = filedialog.asksaveasfilename(
                        initialfile=file_name,
                        defaultextension=".pdf",
                        filetypes=[("PDF datoteke", "*.pdf")],
                        title="Shrani izpisek kot..."
                    )
                    root.destroy()
                    
                    if saved_path:
                        with open(saved_path, "wb") as f:
                            f.write(pdf_bytes)
                        st.success(f"✅ Izpisek uspešno prenesen in shranjen:\n`{saved_path}`")
                    else:
                        st.warning("Prenos preklican.")
            except EzkError as e:
                st.error(str(e))
            except Exception as e:
                st.error(f"Nepričakovana napaka med pridobivanjem izpiska: {e}")

    if st.session_state.ezk_last_pdf:
        fname, fbytes = st.session_state.ezk_last_pdf
        st.download_button(
            "⬇️ Prenesi zadnji izpisek (.pdf)", data=fbytes, file_name=fname,
            mime="application/pdf", use_container_width=True,
        )
     # --------------------------------------------------------------------------
    # SAMODEJNO SERIJSKO ISKANJE (EXCEL + PDF + SLIKA)
    # --------------------------------------------------------------------------
    st.divider()
    st.subheader("📑 Serijsko iskanje iz seznama (Excel / PDF / Slika)")
    st.caption(
        "Naložite datoteko s seznamom parcel. Sistem jo bo analiziral in "
        "samodejno pridobil vse izpiske iz eSodstva v zavihek za pregled lastnikov."
    )

    batch_upload = st.file_uploader(
        "Naloži datoteko (.xlsm, .xlsx, .xls, .pdf, .png, .jpg, .jpeg)", 
        type=["xlsm", "xlsx", "xls", "pdf", "png", "jpg", "jpeg"],
        key="batch_uploader"
    )

    if batch_upload:
        parcels_to_fetch = []

        # 1. OBNAVA IN BRANJE EXCEL DATOTEKE (.xlsm, .xlsx, .xls)
        if batch_upload.name.lower().endswith(('.xlsm', '.xlsx', '.xls')):
            try:
                df = pd.read_excel(batch_upload, engine='openpyxl') 
                st.markdown("**Predogled prebranih stolpcev:**")
                st.dataframe(df.head(3))

                st.markdown("**1. Povežite stolpce z ustreznimi podatki:**")
                c1, c2 = st.columns(2)
                with c1:
                    ko_col = st.selectbox("Kateri stolpec vsebuje šifro Katastrske občine?", options=["-- Izberi --"] + list(df.columns))
                with c2:
                    parc_col = st.selectbox("Kateri stolpec vsebuje Parcelno številko?", options=["-- Izberi --"] + list(df.columns))

                if ko_col != "-- Izberi --" and parc_col != "-- Izberi --":
                    for _, row in df.iterrows():
                        k_val = str(row[ko_col]).strip()
                        p_val = str(row[parc_col]).strip()
                        
                        if k_val and p_val and k_val.lower() != 'nan' and p_val.lower() != 'nan':
                            if k_val.endswith('.0'): k_val = k_val[:-2]
                            if p_val.endswith('.0'): p_val = p_val[:-2]
                            parcels_to_fetch.append((k_val, p_val))
                            
            except Exception as e:
                st.error(f"Napaka pri obdelavi Excel datoteke: {e}")

        # 2. OBNAVA IN BRANJE PDF DATOTEKE
        elif batch_upload.name.lower().endswith('.pdf'):
            try:
                with st.spinner("Berem in prepoznavam parcele iz PDF datoteke..."):
                    pdf_text = extract_pdf_text(batch_upload.read())
                    detected_pairs = parse_ko_parcels_from_text(pdf_text)
                    
                    if detected_pairs:
                        st.success(f"V PDF datoteki sem uspel prepoznati {len(detected_pairs)} parcelnih vnosov!")
                    else:
                        st.warning("V PDF datoteki nisem uspel prepoznati nobenih parcel.")

                    # NOVO: urejljiva tabela - dopolnite/popravite ročno, če je treba
                    df_detected = pd.DataFrame(detected_pairs, columns=["Šifra KO", "Parcelna številka"])
                    df_edited = st.data_editor(
                        df_detected, num_rows="dynamic", use_container_width=True,
                        key="batch_pdf_editor",
                    )
                    parcels_to_fetch = [
                        (str(r["Šifra KO"]).strip(), str(r["Parcelna številka"]).strip())
                        for _, r in df_edited.iterrows()
                        if str(r["Šifra KO"]).strip() and str(r["Parcelna številka"]).strip()
                    ]
            except Exception as e:
                st.error(f"Napaka pri branju PDF datoteke: {e}")

       # 3. OBNAVA IN BRANJE SLIKE (.png, .jpg, .jpeg)
        elif batch_upload.name.lower().endswith(('.png', '.jpg', '.jpeg')):
            try:
                with st.spinner("Berem in prepoznavam besedilo na sliki (OCR)..."):
                    import pytesseract
                    from PIL import Image, ImageOps

                    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                    
                    img = Image.open(batch_upload)

                    # NOVO: predobdelava slike izboljša prepoznavo znakov pri fotografijah
                    # s telefona (siva lestvica, avtomatski kontrast, povečava manjših slik).
                    proc_img = ImageOps.grayscale(img)
                    proc_img = ImageOps.autocontrast(proc_img, cutoff=1)
                    pw, ph = proc_img.size
                    if max(pw, ph) < 3000:
                        proc_img = proc_img.resize((int(pw * 1.8), int(ph * 1.8)), Image.LANCZOS)

                    # KLJUČNO: dodan config='--psm 6' prisili bralnik, da ne meša stolpcev in bere strogo po vrsticah!
                    image_text = pytesseract.image_to_string(proc_img, lang='slv+eng', config='--psm 6')
                    
                    # Uporabimo izboljšano funkcijo za iskanje KO in parcel
                    detected_pairs = parse_ko_parcels_from_text(image_text)
                    
                    if detected_pairs:
                        st.success(f"Na sliki sem uspel prepoznati {len(detected_pairs)} parcelnih vnosov!")
                    else:
                        st.warning("Na sliki nisem uspel prepoznati nobenih parcel. Poskusite s sliko v višji ločljivosti in preverite, da so podatki čitljivi.")

                    # NOVO: tabela je zdaj UREJLJIVA - OCR ni nikoli 100% zanesljiv (npr. zabrisane
                    # ali obledele vrstice), zato lahko tukaj ročno popravite/dodate/izbrišete vrstico,
                    # preden zaženete prenos. Primerjajte s sliko in dopolnite manjkajoče vnose.
                    df_detected = pd.DataFrame(detected_pairs, columns=["Šifra KO", "Parcelna številka"])
                    df_edited = st.data_editor(
                        df_detected, num_rows="dynamic", use_container_width=True,
                        key="batch_image_editor",
                    )
                    parcels_to_fetch = [
                        (str(r["Šifra KO"]).strip(), str(r["Parcelna številka"]).strip())
                        for _, r in df_edited.iterrows()
                        if str(r["Šifra KO"]).strip() and str(r["Parcelna številka"]).strip()
                    ]
            except ImportError:
                st.error("Za branje iz slike manjkata knjižnici `pytesseract` in `Pillow`. Namestite ju preko terminala s: `pip install pytesseract pillow`.")
            except pytesseract.TesseractNotFoundError:
                st.error("Na sistemu ni nameščen Tesseract OCR. Namestite ga na sistem preden poskusite prebrati sliko.")
            except Exception as e:
                st.error(f"Napaka pri branju slike: {e}")

        # ZAGON MASOVNEGA PRENOSA V ZEMLJIŠKI KNJIGI
        if parcels_to_fetch:
            unique_parcels = list(dict.fromkeys(parcels_to_fetch))

            st.markdown("**2. Potrditev:**")
            if st.button(f"▶️ ZAŽENI PRENOS ZA {len(unique_parcels)} PARCEL", type="primary", use_container_width=True):
                progress_bar = st.progress(0.0)
                status_text = st.empty()
                success_count = 0
                
                for i, (k_val, p_val) in enumerate(unique_parcels):
                    status_text.text(f"Prenašam ({i+1}/{len(unique_parcels)}): KO {k_val}, Parc. št. {p_val} ...")
                    try:
                        pdf_bytes = cached_fetch_ezk_pdf(
                            _state=st.session_state.ezk_storage_state,
                            katastrska_obcina=k_val,
                            parcelna_stevilka=p_val,
                            headless=not st.session_state.ezk_debug_mode,
                        )
                        
                        file_name = f"ezk_{k_val}_{p_val}.pdf"
                        file_name = _safe_filename_part(file_name.replace(".pdf", "")) + ".pdf"
                        file_id = f"{file_name}_{len(pdf_bytes)}"
                        
                        existing_ids = {r["file_id"] for r in st.session_state.bulk_pdf_records}
                        
                        if file_id not in existing_ids:
                            if len(st.session_state.bulk_pdf_records) < MAX_BULK_PDFS:
                                st.session_state.bulk_pdf_records.append(
                                    {"file_id": file_id, "name": file_name, "text": extract_pdf_text(pdf_bytes)}
                                )
                                success_count += 1
                            else:
                                st.warning(f"Dosežen je maksimum {MAX_BULK_PDFS} hkrati naloženih izpiskov!")
                                break
                    except Exception as e:
                        st.error(f"Napaka pri prenosu (KO: {k_val}, Parc: {p_val}): {e}")
                        
                    progress_bar.progress((i + 1) / len(unique_parcels))
                
                status_text.success(
                    f"✅ Serijski prenos je končan! Uspešno prenesenih in dodanih {success_count} "
                    f"izpiskov. Rezultate si lahko pogledate v zavihku **{TAB_BULK_LABEL}**."
                )

with tab_projekt:
    st.subheader("🏗️ Podatki o projektu (gradnji)")
    st.caption(
        "Vnesite osnovne podatke o projektu/gradnji. Ti podatki se pri izpolnjevanju "
        "pogodbe samodejno vpišejo v tabelo na začetku pogodbe (Naziv gradnje, "
        "Št. dokumentacije, Datum izdelave)."
    )

    with st.form("projekt_form"):
        naziv_gradnje = st.text_input(
            "Naziv gradnje:",
            value=st.session_state.projekt_info.get("naziv_gradnje", ""),
            placeholder="npr. IZGRADNJA NOVEGA NN IZVODA ZA NASELJE OPPN Š 18",
        )
        st_dokumentacije = st.text_input(
            "Št. dokumentacije:",
            value=st.session_state.projekt_info.get("st_dokumentacije", ""),
            placeholder="npr. Popis del z oceno stroškov",
        )
        datum_izdelave = st.text_input(
            "Datum izdelave:",
            value=st.session_state.projekt_info.get("datum_izdelave", ""),
            placeholder="npr. 2026",
        )
        projekt_submitted = st.form_submit_button("💾 Shrani podatke o projektu", type="primary", use_container_width=True)

    if projekt_submitted:
        st.session_state.projekt_info = {
            "naziv_gradnje": naziv_gradnje.strip(),
            "st_dokumentacije": st_dokumentacije.strip(),
            "datum_izdelave": datum_izdelave.strip(),
        }

        # Če je predloga že naložena in ima rdeča polja v tabeli PROJEKT,
        # takoj posodobimo tudi trenutno stanje dokumenta v zavihku
        # "Izpolnjevanje pogodbe" (če uporabnik še ni prenesel podatkov o lastniku).
        if st.session_state.doc is not None and st.session_state.fields:
            n_applied = apply_projekt_fields_to_doc(
                st.session_state.doc, st.session_state.fields, st.session_state.projekt_info
            )
            if n_applied:
                st.session_state.print_pdf_bytes = None

        st.success("Podatki o projektu so shranjeni in bodo samodejno vpisani v pogodbo.")

    if any(st.session_state.projekt_info.values()):
        st.divider()
        st.markdown("**Trenutno shranjeni podatki:**")
        for key, label in PROJEKT_LABELS_SL.items():
            val = st.session_state.projekt_info.get(key, "")
            st.markdown(f"- **{label}:** {val or '_ni vneseno_'}")

with tab_settings:
    if "settings_authenticated" not in st.session_state:
        st.session_state.settings_authenticated = False

    if not st.session_state.settings_authenticated:
        st.subheader("🔒 Nastavitve so zaklenjene")
        st.caption("Do nastavitev lahko dostopa samo administrator.")
        with st.form("settings_login_form"):
            _login_user = st.text_input("Uporabniško ime")
            _login_pass = st.text_input("Geslo", type="password")
            _login_submit = st.form_submit_button("Prijava")
        if _login_submit:
            if _login_user == "admin" and _login_pass == "admin":
                st.session_state.settings_authenticated = True
                st.rerun()
            else:
                st.error("Napačno uporabniško ime ali geslo.")
    else:
        _logout_col1, _logout_col2 = st.columns([5, 1])
        with _logout_col2:
            if st.button("🚪 Odjava", key="settings_logout"):
                st.session_state.settings_authenticated = False
                st.rerun()

        st.subheader("⚙️ Nastavitve API strežnika in modela")
        st.caption("Prilagodite parametre za komunikacijo z jezikovnimi modeli in upravljajte ključe.")

        _default_settings = _load_default_settings()
        _default_provider = _default_settings.get("provider")

        nvidia_key_env = (os.environ.get("NVIDIA_API_KEY") or "").strip()
        gemini_key_env = (os.environ.get("GEMINI_API_KEY") or "").strip()

        env_key_options = []
        if nvidia_key_env:
            env_key_options.append({
                "label": "NVIDIA API ključ (.env)",
                "provider": PROVIDER_NVIDIA,
                "key": nvidia_key_env,
                "base_url": (os.environ.get("NVIDIA_BASE_URL") or "").strip() or DEFAULT_NVIDIA_BASE_URL,
            })
        if gemini_key_env:
            env_key_options.append({
                "label": "Gemini API ključ (.env)",
                "provider": PROVIDER_GEMINI,
                "key": gemini_key_env,
                "base_url": (os.environ.get("GEMINI_BASE_URL") or "").strip() or DEFAULT_GEMINI_BASE_URL,
            })

        override_key_input = ""
        desired_base_url = ""
        provider_signal = ""
        new_key_entered = False

        if len(env_key_options) >= 2:
            st.markdown(
                '<span class="status-pill pill-ok">✓ Najdena 2 API ključa (.env / okolje)</span>',
                unsafe_allow_html=True,
            )
        
            default_index = 0
            if _default_provider:
                for i, opt in enumerate(env_key_options):
                    if opt["provider"] == _default_provider:
                        default_index = i
                        break
        
            chosen_label = st.radio(
                "Kateri ključ naj se uporabi?",
                options=[o["label"] for o in env_key_options],
                index=default_index,
                horizontal=True,
                key="env_key_choice",
            )
            chosen = next(o for o in env_key_options if o["label"] == chosen_label)
            api_key_input = chosen["key"]
            desired_base_url = chosen["base_url"]
            provider_signal = f"env:{chosen_label}"
        
            if st.button("💾 Shrani izbiro ključa kot privzeto"):
                _write_default_settings(provider=chosen["provider"])
                st.success(f"Izbira ključa ({chosen_label}) bo pri naslednjem zagonu aplikacije privzeta.")

            with st.expander("Uporabi drug ključ za to sejo"):
                override_key_input = st.text_input(
                    "API ključ (začasna prepoznava)", type="password", key="session_override_key",
                    help="Vnesite drug NVIDIA ali Gemini ključ samo za to sejo.",
                )
                if override_key_input:
                    api_key_input = override_key_input
                    desired_base_url = base_url_for_provider(detect_provider(api_key=override_key_input))
                    provider_signal = f"override:{override_key_input}"
                    new_key_entered = True
        elif len(env_key_options) == 1:
            chosen = env_key_options[0]
            st.markdown(
                f'<span class="status-pill pill-ok">✓ {chosen["label"].replace(" (.env)", "")} najden (.env / okolje)</span>',
                unsafe_allow_html=True,
            )
            api_key_input = chosen["key"]
            desired_base_url = chosen["base_url"]
            provider_signal = f"env:{chosen['label']}"
            with st.expander("Uporabi drug ključ za to sejo"):
                override_key_input = st.text_input(
                    "API ključ (začasna prepoznava)", type="password", key="session_override_key",
                )
                if override_key_input:
                    api_key_input = override_key_input
                    desired_base_url = base_url_for_provider(detect_provider(api_key=override_key_input))
                    provider_signal = f"override:{override_key_input}"
                    new_key_entered = True
        else:
            api_key_input = st.text_input(
                "API ključ", type="password",
                help="NVIDIA ali Gemini API ključ. Shrani v .env spodaj.",
            )
            if api_key_input:
                desired_base_url = base_url_for_provider(detect_provider(api_key=api_key_input))
                new_key_entered = True
            provider_signal = f"manual:{api_key_input}"

        _default_base_url = (_default_settings.get("base_url") or "").strip()
        _default_model = (_default_settings.get("model") or "").strip()

        if "base_url_input" not in st.session_state:
            st.session_state.base_url_input = _default_base_url or desired_base_url or DEFAULT_NVIDIA_BASE_URL
            st.session_state.last_provider_signal = provider_signal
        elif provider_signal != st.session_state.get("last_provider_signal"):
            if desired_base_url:
                st.session_state.base_url_input = desired_base_url
            st.session_state.last_provider_signal = provider_signal

        active_provider = detect_provider(api_key=api_key_input, base_url=st.session_state.base_url_input)
        provider_label = PROVIDER_LABELS[active_provider]

        base_url_input = st.text_input(
            f"Naslov {provider_label} (base URL)",
            key="base_url_input",
            placeholder=base_url_for_provider(active_provider),
        )

        active_provider = detect_provider(api_key=api_key_input, base_url=base_url_input)
        provider_label = PROVIDER_LABELS[active_provider]

        proxy_input = st.text_input(
            "HTTP(S) proxy (neobvezno)",
            value=os.environ.get("HTTPS_PROXY", os.environ.get("HTTP_PROXY", "")),
            placeholder="http://uporabnik:geslo@proxy.podjetje.si:8080",
        )

        if new_key_entered:
            save_env_col, save_env_status_col = st.columns([1, 3])
            with save_env_col:
                save_to_env_clicked = st.button(
                    "💾 Shrani ključ in naslov v .env", use_container_width=True,
                    disabled=not api_key_input,
                )
            if save_to_env_clicked:
                try:
                    _save_provider_credentials_to_env(active_provider, api_key_input, base_url_input)
                    with save_env_status_col:
                        st.success("Ključ in naslov uspešno shranjena v lokalni `.env`.")
                except Exception as e:
                    with save_env_status_col:
                        st.error(f"Shranjevanje v .env ni uspelo: {e}")

        st.divider()
        st.markdown("**Model**")

        if "available_models" not in st.session_state:
            st.session_state.available_models = []
        if "models_fetched_for" not in st.session_state:
            st.session_state.models_fetched_for = None
        if "models_fetch_error" not in st.session_state:
            st.session_state.models_fetch_error = None

        effective_base_url = (base_url_input or "").strip() or base_url_for_provider(active_provider)
        fetch_signature = (effective_base_url, api_key_input or "")

        refresh_col, status_col = st.columns([1, 3])
        with refresh_col:
            manual_refresh = st.button("🔄 Osveži modele", use_container_width=True)

        if manual_refresh or (api_key_input and st.session_state.models_fetched_for != fetch_signature):
            try:
                with st.spinner("Pridobivam seznam modelov ..."):
                    fetched = fetch_available_models(
                        api_key=api_key_input or None,
                        base_url=base_url_input or None,
                        proxy_url=proxy_input or None,
                    )
                st.session_state.available_models = fetched
                st.session_state.models_fetch_error = None
            except LLMExtractionError as e:
                st.session_state.models_fetch_error = str(e)
            st.session_state.models_fetched_for = fetch_signature

        with status_col:
            if st.session_state.models_fetch_error:
                st.warning(f"Neuspešno pridobivanje modelov: {st.session_state.models_fetch_error}")
            elif st.session_state.available_models:
                st.caption(f"✓ Pridobljenih {len(st.session_state.available_models)} modelov.")
            else:
                st.caption("Seznam modelov bo samodejno pridobljen ob vnosu ključa.")

        fallback_models = PROVIDER_FALLBACK_MODELS.get(active_provider, PROVIDER_FALLBACK_MODELS[PROVIDER_NVIDIA])
        model_options = list(st.session_state.available_models) if st.session_state.available_models else fallback_models
        if _default_model and _default_model not in model_options:
            model_options = [_default_model] + model_options
        model_options = model_options + ["Drug model (vnesi spodaj) ..."]

        default_model_index = model_options.index(_default_model) if _default_model in model_options else 0

        model_select_col, model_default_col = st.columns([3, 1])
        with model_select_col:
            model_choice = st.selectbox(
                f"Model ({provider_label})",
                options=model_options,
                index=default_model_index,
            )
        with model_default_col:
            st.write("")
            st.write("")
            set_default_clicked = st.button(
                "⭐ Nastavi kot privzet", use_container_width=True,
            )
        if model_choice == "Drug model (vnesi spodaj) ...":
            model_choice = st.text_input(
                f"Model slug ({provider_label})", value=_default_model or "", placeholder="npr. meta/llama-3.3-70b-instruct",
            ).strip() or DEFAULT_MODEL

        # Shrani trenutno veljavne nastavitve v session_state, da jih lahko
        # uporabijo tudi drugi zavihki (npr. skupinski pregled), tudi ko je
        # zavihek "Nastavitve" zaklenjen in se ta blok ne izvede.
        st.session_state.current_api_key = api_key_input
        st.session_state.current_proxy_url = proxy_input
        st.session_state.current_base_url = base_url_input
        st.session_state.current_model = model_choice

        if set_default_clicked:
            _write_default_settings(effective_base_url, model_choice, active_provider)
            st.success(
                f"Model '{model_choice}' in naslov '{effective_base_url}' sta nastavljena kot privzeta - "
                "ostaneta enaka tudi ob naslednjem zagonu programa."
            )
            st.rerun()

        st.divider()
        if st.button(f"🔧 Preizkusi povezavo do {provider_label}", use_container_width=True):
            with st.spinner("Preverjam omrežno povezavo ..."):
                result = test_network_connectivity(proxy_url=proxy_input or None, base_url=base_url_input or None)
            (st.success if result.startswith("✅") else st.error)(result)

with tab_fill:
    # ----------------------------------------------------------------------
    # TRIJE STOLPCI
    # ----------------------------------------------------------------------
    left, mid, right = st.columns([1, 1.5, 1.4], gap="medium")

    # ============================== LEVI STOLPEC ==============================
    with left:
        st.subheader("📂 POGODBA O USTANOVITVI SLUŽNOSTNE PRAVICE")

        auto_info = st.session_state.get("auto_template_info")
        if auto_info and st.session_state.template_file_id == auto_info.get("file_id"):
            st.success(f"✅ Samodejno naložena predloga: **{auto_info['label']}**")
            if auto_info.get("gender_uncertain"):
                st.caption(
                    "ℹ️ Spol lastnika/lastnice ni bil eksplicitno zaznan v podatkih - "
                    "izbran na podlagi imena. Preverite, ali je predloga pravilna "
                    "(gl. spodaj 'Ročno naloži drugo predlogo')."
                )
        elif st.session_state.doc is None:
            st.info(
                "Predloga se samodejno naloži, ko v zavihku "
                "'📊 PREGLED LASTNIKOV PARCEL IZ ZEMLJIŠKE KNJIGE' kliknete "
                "'📤 PRENESI V IZPOLNJEVANJE POGODBE' - glede na število "
                "lastnikov (in pri 1 lastniku glede na spol) se naloži "
                "pravilna različica pogodbe."
            )

        with st.expander("🔧 Ročno naloži drugo predlogo (.docx)", expanded=False):
            docx_upload = st.file_uploader("Predloga pogodbe (.docx)", type=["docx"], key="docx_uploader")

            if docx_upload is not None:
                file_id = f"{docx_upload.name}_{docx_upload.size}"
                if file_id != st.session_state.template_file_id:
                    _activate_template(docx_upload.read(), file_id)
                    st.session_state.auto_template_info = None
                    st.rerun()
                    
                st.divider()
                st.markdown("**💾 Dodaj predlogo ali posodobi obstoječo predlogo**")
                st.caption("Shranite to predlogo v mapo `templates/`, da se bo v prihodnje izbrala samodejno.")
                save_name = st.text_input("Ime datoteke za shranjevanje:", value=docx_upload.name)
                if st.button("Shrani predlogo", use_container_width=True):
                    try:
                        save_path = TEMPLATES_DIR / save_name
                        with open(save_path, "wb") as f:
                            docx_upload.seek(0)
                            f.write(docx_upload.read())
                        st.success(f"Predloga uspešno shranjena: {save_name}")
                    except Exception as e:
                        st.error(f"Napaka pri shranjevanju: {e}")

        if st.session_state.doc is not None:
            st.markdown(
                f'<span class="status-pill pill-ok">✓ Predloga – {len(st.session_state.fields)} rdečih polj</span>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown('<span class="status-pill pill-neutral">Predloga še ni naložena</span>', unsafe_allow_html=True)

        st.divider()
        if st.button("🔄 Ponastavi vse", use_container_width=True):
            reset_all()
            st.rerun()

    # ============================== SREDNJI STOLPEC ==============================
    with mid:
        if st.session_state.transferred_parcels:
            st.subheader("📥 Preneseni podatki")
            st.markdown(
                f"**{st.session_state.transferred_owner_label}** "
                f"— {len(st.session_state.transferred_parcels)} parcel(e)"
            )
            st.caption(
                "Podatki so bili preneseni iz zavihka 'Skupinski pregled lastnikov' - že izluščeni, "
                "brez novega klica AI modela."
            )
            transfer_disabled = st.session_state.doc is None
            if transfer_disabled:
                st.caption("Najprej v levem stolpcu naložite .docx predlogo.")
            tcol1, tcol2 = st.columns([2, 1])
            with tcol1:
                use_transferred = st.button(
                    "📥 Uporabi prenesene podatke in izpolni predlogo",
                    type="primary", use_container_width=True, disabled=transfer_disabled,
                )
            with tcol2:
                if st.button("✕", key="clear_transferred", help="Počisti prenesene podatke", use_container_width=True):
                    st.session_state.transferred_parcels = None
                    st.session_state.transferred_owner_label = None
                    st.session_state.transferred_zupan = None
                    st.session_state.transferred_maticna = None
                    st.rerun()

            if use_transferred:
                try:
                    parcels = st.session_state.transferred_parcels
                    fresh_doc = load_docx(st.session_state.original_docx_bytes)
                    fresh_fields = parse_docx_fields(fresh_doc, red_hex=st.session_state.red_hex_codes)

                    # --- NOVO: SAMODEJNI VNOS PODATKOV O PROJEKTU (zavihek PROJEKT) ---
                    n_projekt_applied = apply_projekt_fields_to_doc(
                        fresh_doc, fresh_fields, st.session_state.projekt_info
                    )
                    # -------------------------------------------------------------------

                    # --- NOVO: UPORABA PRENESENIH PODATKOV O ŽUPANU IN MATIČNI ---
                    zupan = st.session_state.get("transferred_zupan")
                    maticna = st.session_state.get("transferred_maticna")
                    
                    if zupan and maticna:
                        gender = guess_owner_gender({"ime_priimek": zupan})
                        
                        fields_to_remove = []
                        # 1. Zamenjaj v rdečih poljih (fields) - iščemo "župan" in 10-mestne številke
                        for fid, f in fresh_fields.items():
                            val = f.value.strip()
                            val_lower = val.lower()
                            matched_special = False
                            
                            # Župan
                            if "{IME_ZUPANA}" in val:
                                apply_field_value(f, f.value.replace("{IME_ZUPANA}", zupan))
                                matched_special = True
                            elif "župan" in val_lower and "občina" not in val_lower and "obcina" not in val_lower:
                                is_cap = val.startswith("Ž") or val.startswith("ž") == False
                                if gender == "Z":
                                    naziv = "Županja" if is_cap else "županja"
                                else:
                                    naziv = "Župan" if is_cap else "župan"
                                apply_field_value(f, f"{naziv} {zupan}")
                                matched_special = True
                                
                            # Matična številka
                            if "{MATICNA_OBCINE}" in val:
                                apply_field_value(f, f.value.replace("{MATICNA_OBCINE}", maticna))
                                matched_special = True
                            elif re.fullmatch(r'\s*\d{10}\s*', val):
                                apply_field_value(f, maticna)
                                matched_special = True
                                
                            # Če smo polje zamenjali, ga odstranimo iz fresh_fields, 
                            # da ga AI ne bo znova povozil in spremenil nazaj v ime občine!
                            if matched_special:
                                fields_to_remove.append(fid)
                                
                        for fid in fields_to_remove:
                            del fresh_fields[fid]
                                
                        # 2. Zamenjava statičnih oznak v preostalem navadnem besedilu
                        zamenjave_navadne = {
                            "{IME_ZUPANA}": zupan,
                            "{MATICNA_OBCINE}": maticna
                        }
                        def replace_in_p(p):
                            for run in p.runs:
                                for k, v in zamenjave_navadne.items():
                                    if k in run.text:
                                        run.text = run.text.replace(k, v)
                                        
                        for p in fresh_doc.paragraphs:
                            replace_in_p(p)
                        for table in fresh_doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        replace_in_p(p)
                    # -----------------------------------------------------

                    report = match_and_apply(fresh_fields, parcels[0])

                    unique_owners = dedupe_owners(parcels)
                    # --- NOVO: Če je lastnik občina, ga tretirajmo kot edinega lastnika za razširjanje vrstic ---
                    owner_name = parcels[0].get("ime_priimek", "").lower()
                    is_municipality = "občina" in owner_name or "obcina" in owner_name
                    
                    if len(unique_owners) > 1 and not is_municipality:
                        added_owner_rows = expand_owner_rows(fresh_doc, fresh_fields, report, unique_owners)
                        assign_owners_to_slots(fresh_fields, report, unique_owners)
                        expand_free_text_owner_mentions(fresh_fields, report, unique_owners)
                        n_summary_fields = fill_owner_list_summary_fields(fresh_fields, report, unique_owners)
                        st.info(
                            f"Vpisanih {len(unique_owners)} lastnikov v predlogo"
                            + (f" (dodanih {added_owner_rows} novih vrstic)." if added_owner_rows else ".")
                            + (f" Posodobljenih {n_summary_fields} povzetnih polj s seznamom vseh lastnikov." if n_summary_fields else "")
                        )
                    elif is_municipality:
                        # Če je občina, polja samo uskladimo s prvo (edino) enoto, brez razširjanja
                        assign_owners_to_slots(fresh_fields, report, [unique_owners[0]])
                    # --------------------------------------------------------------------------

                    unique_parcels = dedupe_parcels(parcels)
                    if len(unique_parcels) > 1:
                        n_merged = merge_parcels_single_row(fresh_fields, report, unique_parcels)
                        st.info(
                            f"Vseh {len(unique_parcels)} parcel je zapisanih v ISTO vrstico "
                            f"(parcelne številke: {', '.join(p.get('parcelna_stevilka', '') for p in unique_parcels)})."
                        )
                    else:
                        collapse_duplicate_parcel_cell_paragraphs(fresh_fields, report, unique_parcels)

                    st.session_state.doc = fresh_doc
                    st.session_state.fields = fresh_fields
                    st.session_state.match_report = report
                    st.session_state.parcels_data = parcels
                    st.session_state.print_pdf_bytes = None
                    
                    merged_changes: dict[str, dict] = {}
                    change_order: list[str] = []
                    for r in report:
                        if r.field_id not in fresh_fields:
                            continue
                        if r.field_id not in merged_changes:
                            merged_changes[r.field_id] = {"role": r.role, "old_value": r.old_value}
                            change_order.append(r.field_id)
                        else:
                            merged_changes[r.field_id]["role"] = r.role

                    field_changes = {}
                    for fid in change_order:
                        old_value = merged_changes[fid]["old_value"]
                        new_value = _current_field_text(fresh_fields[fid])
                        if old_value.strip() != new_value.strip():
                            field_changes[fid] = {
                                "role": merged_changes[fid]["role"],
                                "old_value": old_value,
                                "new_value": new_value,
                            }
                    st.session_state.field_changes = field_changes
                    st.session_state.field_roles = {fid: c["role"] for fid, c in field_changes.items()}

                    n_changed = sum(1 for r in report if r.changed)
                    st.success(
                        f"Najdenih {len(report)} ujemajočih polj za primarno parcelo, "
                        f"dejansko spremenjenih: {n_changed}."
                        + (f" Vpisanih {n_projekt_applied} podatkov o projektu." if n_projekt_applied else "")
                    )
                except Exception as e:
                    st.error(f"Napaka pri izpolnjevanju predloge: {e}")
            st.divider()

        if st.session_state.doc is None:
            st.info("Najprej v levem stolpcu naložite .docx predlogo.")
        elif st.session_state.match_report is None:
            st.info(
                "Najprej prenesite podatke lastnika (zavihek 'Skupinski pregled') in kliknite "
                "'📥 Uporabi prenesene podatke in izpolni predlogo'."
            )
        else:
            st.subheader("📝 Podatki iz zemljiške knjige")

            if st.session_state.parcels_data:
                for i, parcel in enumerate(st.session_state.parcels_data):
                    with st.container(border=True):
                        st.markdown(f"**Podatki o lastniku {i + 1}**")
                        
                        for key in LAND_REGISTRY_KEYS:
                            val = parcel.get(key, "")
                            st.markdown(f"**{LABELS_SL.get(key, key)}:** {val or '_ni najdeno_'}")
                            
                            # --- NOVO: Prikaz matične in župana takoj pod naslovom ---
                            if key == "naslov" and i == 0 and st.session_state.get("transferred_zupan"):
                                st.markdown(f"**Matična št.:** {st.session_state.get('transferred_maticna')}")
                                st.markdown(f"**Župan/ja:** {st.session_state.get('transferred_zupan')}")
                            # -------------------------------------------------------

            st.subheader("📝 Rdeča polja v predlogi")
            if not st.session_state.fields:
                st.info("Ni najdenih rdečih polj v naloženi predlogi.")
            else:
                role_map = st.session_state.field_roles or {}

                matched_fields = [f for fid, f in st.session_state.fields.items() if fid in role_map]
                other_fields = [f for fid, f in st.session_state.fields.items() if fid not in role_map]

                if matched_fields:
                    with st.expander(
                        f"**✓ SPREMENJENA POLJA V POGODBI O USTANOVITVI SLUŽNOSTNE PRAVICE** ({len(matched_fields)})",
                        expanded=False, 
                    ):
                        changes = st.session_state.field_changes or {}
                        for f in matched_fields:
                            role = role_map[f.id]
                            st.markdown(
                                f"{LABELS_SL.get(role, role)} "
                                '<span class="status-pill pill-ok">spremenjeno</span>',
                                unsafe_allow_html=True,
                            )
                            change = changes.get(f.id)
                            if change:
                                st.caption(f"prej: {change['old_value'].strip()}  →  zdaj: {change['new_value'].strip()}")
                            st.text_input("vrednost", value=_current_field_text(f), key=f"field_input_{f.id}", label_visibility="collapsed")
                            if len(f.locations) > 1:
                                st.markdown(f'<div class="field-caption">pojavi se {len(f.locations)}× v dokumentu</div>', unsafe_allow_html=True)

                if other_fields:
                    with st.expander(f"**Ostala rdeča polja (ročni vnos)** ({len(other_fields)})", expanded=True):
                        for f in other_fields:
                            st.text_input(f.label, value=_current_field_text(f), key=f"field_input_{f.id}")
                            if f.is_blank_placeholder:
                                st.markdown('<div class="field-caption">prazno polje v predlogi</div>', unsafe_allow_html=True)

                sync_fields_from_widgets()

    # ============================== DESNI STOLPEC ==============================
    with right:
        st.subheader("👁️ Predogled")
        if st.session_state.doc is None:
            st.markdown("<p style='color:#9CA3AF'>Predogled bo prikazan po nalaganju predloge.</p>", unsafe_allow_html=True)
        else:
            if st.button("🔄 Osveži predogled", use_container_width=True):
                sync_fields_from_widgets()
            
            preview_rendered = False
            
            # Najprej poskusi najboljši (PDF) predogled preko LibreOffice
            if preview_engine.is_available():
                try:
                    with st.spinner("Pripravljam natančen PDF predogled..."):
                        current_docx_bytes = save_docx(st.session_state.doc)
                        pdf_bytes = preview_engine.docx_to_pdf_bytes(current_docx_bytes)
                        pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")
                        
                        st.markdown(
                            f'<iframe src="data:application/pdf;base64,{pdf_b64}#toolbar=0&navpanes=0" '
                            f'width="100%" height="800px" style="border: 1px solid var(--border-gray); border-radius: 8px;"></iframe>', 
                            unsafe_allow_html=True
                        )
                        preview_rendered = True
                except Exception as e:
                    st.error(f"Napaka pri pripravi PDF predogleda: {e}")
            
            # Če PDF metoda ni uspela ali ni podprta, uporabi poenostavljen mammoth
            if not preview_rendered:
                if not MAMMOTH_AVAILABLE:
                    st.warning("Za natančen predogled priporočamo namestitev LibreOffice (soffice). Alternativno namestite knjižnico 'mammoth'.")
                else:
                    try:
                        import io
                        with st.spinner("Pripravljam poenostavljen HTML predogled (mammoth)..."):
                            current_docx_bytes = save_docx(st.session_state.doc)
                            buf = io.BytesIO(current_docx_bytes)
                            result = mammoth.convert_to_html(buf)
                            html_content = result.value
                            
                            st.markdown(
                                f'<div class="preview-frame" style="background: white; color: black; padding: 20px; border-radius: 8px; border: 1px solid var(--border-gray); max-height: 800px; overflow-y: auto;">{html_content}</div>', 
                                unsafe_allow_html=True
                            )
                    except Exception as e:
                        st.error(f"Napaka pri pripravi HTML predogleda: {e}")

        if st.session_state.doc is not None:
            st.subheader("🖨️ Tiskanje")
            if not preview_engine.is_available():
                st.button("🖨️ Natisni dokument", disabled=True, use_container_width=True)
                st.caption("Za neposredno tiskanje je potreben LibreOffice (glej README.md).")
            else:
                if st.button("🖨️ Natisni dokument (predogled)", use_container_width=True):
                    sync_fields_from_widgets()
                    try:
                        with st.spinner("Pripravljam PDF za tiskanje ..."):
                            print_bytes = save_docx(st.session_state.doc)
                            pdf_bytes = preview_engine.docx_to_pdf_bytes(print_bytes)
                        st.session_state.print_pdf_bytes = pdf_bytes
                    except Exception as e:
                        st.error(f"Napaka pri pripravi tiskanja: {e}")

                if st.session_state.get("print_pdf_bytes"):
                    pdf_b64 = base64.b64encode(st.session_state.print_pdf_bytes).decode("utf-8")
                    components.html(
                        f"""
                        <div>
                            <button id="print-btn" style="width:100%;padding:10px;border-radius:8px;
                                border:1px solid #1E4A7B;background:#1E4A7B;color:white;font-weight:600;
                                cursor:pointer;margin-bottom:10px;">🖨️ Odpri predogled in natisni</button>
                            <iframe id="print-frame" style="display:none;"
                                src="data:application/pdf;base64,{pdf_b64}"></iframe>
                        </div>
                        <script>
                        document.getElementById('print-btn').addEventListener('click', function() {{
                            var w = window.open("data:application/pdf;base64,{pdf_b64}", "_blank");
                            if (w) {{
                                w.addEventListener('load', function() {{
                                    try {{ w.print(); }} catch (e) {{}}
                                }});
                            }}
                        }});
                        </script>
                        """,
                        height=60,
                    )

            st.subheader("📥 Izvoz")
            sync_fields_from_widgets()
            
            # --- NOVO: Preveri, če je "XXX" še vedno prisoten v katerem izmed polj ---
            has_xxx = any("XXX" in _current_field_text(f) for f in st.session_state.fields.values())
            
            if has_xxx:
                st.error("NE MOREM SHRANITI! MANJKA ŠTEVILKA POGODBE (besedilo 'XXX' še ni zamenjano)!")
                st.button("💾 SHRANI POGODBO", disabled=True, use_container_width=True, key="btn_shrani_disabled")
            else:
                import docx
                from docx.shared import RGBColor
                
                # Začasno shrani originalne barve in jih obarvaj črno za izvoz
                original_colors = []
                for f in st.session_state.fields.values():
                    for loc in f.locations:
                        for run in loc.runs:
                            if run.font and run.font.color:
                                orig_color = run.font.color.rgb
                                original_colors.append((run, orig_color))
                                run.font.color.rgb = RGBColor(0, 0, 0)
                
                export_bytes = save_docx(st.session_state.doc)
                
                # Obnovi originalne barve za prikaz predogleda v aplikaciji
                for run, orig_color in original_colors:
                    if orig_color:
                        run.font.color.rgb = orig_color
            
                # Prikaži samo gumb za shranjevanje preko poljubne mape
                if st.button("💾 SHRANI POGODBO", type="primary", use_container_width=True):
                    import tkinter as tk
                    from tkinter import filedialog
                    
                    root = tk.Tk()
                    root.withdraw()
                    root.wm_attributes('-topmost', 1)
                    
                    saved_path = filedialog.asksaveasfilename(
                        initialfile=build_export_filename(),
                        defaultextension=".docx",
                        filetypes=[("Word dokument", "*.docx")],
                        title="Shrani izpolnjeno pogodbo kot..."
                    )
                    root.destroy()
                    
                    if saved_path:
                        with open(saved_path, "wb") as f:
                            f.write(export_bytes)
                        st.success(f"✅ Pogodba uspešno shranjena:\n`{saved_path}`")
                    else:
                        st.warning("Shranjevanje preklicano.")

with tab_bulk:
    st.subheader("Naložite PDF izpiske iz zemljiške knjige")
    st.caption(
        f"AI prebere "
        "vsak izpisek posebej, sistem pa jih samodejno razvrsti po lastnikih (in po katastrskih "
        "občinah) - s klikom na ime lastnika ali občino nato vidite vse pripadajoče parcele in deleže."
    )

    bulk_uploads = st.file_uploader(
        "Izpiski iz zemljiške knjige (.pdf)", type=["pdf"], accept_multiple_files=True,
        key=f"bulk_pdf_uploader_{st.session_state.bulk_uploader_version}",
        help=f"Naenkrat lahko naložite največ {MAX_BULK_PDFS} datotek.",
    )

    if bulk_uploads:
        if len(bulk_uploads) > MAX_BULK_PDFS and not st.session_state.bulk_pdf_records:
            st.error(
                f"Naloženih je {len(bulk_uploads)} datotek naenkrat - dovoljenih je največ "
                f"{MAX_BULK_PDFS}. Odstranite nekaj datotek iz izbire."
            )
        else:
            existing_ids = {r["file_id"] for r in st.session_state.bulk_pdf_records}
            changed = False
            limit_hit = False
            for pf in bulk_uploads:
                file_id = f"{pf.name}_{pf.size}"
                if file_id in existing_ids:
                    continue
                if len(st.session_state.bulk_pdf_records) >= MAX_BULK_PDFS:
                    limit_hit = True
                    break
                st.session_state.bulk_pdf_records.append(
                    {"file_id": file_id, "name": pf.name, "text": extract_pdf_text(pf.read())}
                )
                changed = True
            if limit_hit:
                st.warning(f"Doseženih je največ {MAX_BULK_PDFS} datotek - dodatne niso bile dodane.")
            if changed:
                st.rerun()

    if st.session_state.bulk_pdf_records:
        recs = st.session_state.bulk_pdf_records
        with st.expander(f"📄 Naloženi izpiski ({len(recs)}/{MAX_BULK_PDFS})", expanded=False):
            cols_per_row = 3
            for row_start in range(0, len(recs), cols_per_row):
                row_cols = st.columns(cols_per_row)
                for offset, col in enumerate(row_cols):
                    idx = row_start + offset
                    if idx >= len(recs):
                        break
                    rec = recs[idx]
                    with col:
                        c1, c2 = st.columns([4, 1])
                        c1.markdown(f"📄 {rec['name']}")
                        if c2.button("✕", key=f"bulk_remove_{rec['file_id']}", help="Odstrani ta izpisek"):
                            st.session_state.bulk_pdf_records.pop(idx)
                            st.session_state.bulk_extracted = None
                            st.session_state.bulk_owners = None
                            st.session_state.bulk_by_ko = None
                            st.session_state.bulk_uploader_version += 1
                            st.rerun()

        st.divider()

        extract_bulk_disabled = not st.session_state.bulk_pdf_records
        if st.button(
            "🔍 Preglej in razporedi", type="primary", use_container_width=True,
            disabled=extract_bulk_disabled,
        ):
            try:
                extracted_records = []
                undercount_warnings = []
                progress = st.progress(0.0)
                recs = st.session_state.bulk_pdf_records
                for i, rec in enumerate(recs):
                    with st.spinner(f"NVIDIA API analizira ({i + 1}/{len(recs)}): {rec['name']} ..."):
                        data = cached_extract_land_registry_data(
                            pdf_text=rec["text"], api_key=st.session_state.current_api_key or None,
                            proxy_url=st.session_state.current_proxy_url or None, model=st.session_state.current_model,
                            base_url=st.session_state.current_base_url or None,
                        )
                    if isinstance(data, list):
                        owners_found = len(data)
                        for owner_record in data:
                            owner_record = dict(owner_record)
                            owner_record["file_name"] = rec["name"]
                            extracted_records.append(owner_record)
                    else:
                        owners_found = 1
                        data = dict(data)
                        data["file_name"] = rec["name"]
                        extracted_records.append(data)

                    expected_owners = count_probable_owners(rec["text"])
                    if expected_owners > owners_found:
                        undercount_warnings.append(
                            f"⚠️ **{rec['name']}**: besedilo vsebuje znake {expected_owners} "
                            f"lastnikov (\"osebno ime:\"), model pa jih je izluščil samo "
                            f"{owners_found}. Preverite izpisek ročno - morda gre za solastnike, "
                            "ki jih model ni zaznal."
                        )
                    progress.progress((i + 1) / len(recs))
                progress.empty()

                st.session_state.bulk_extracted = extracted_records
                st.session_state.bulk_owners = group_by_ownership_unit(extracted_records)
                st.session_state.bulk_by_ko = group_by_katastrska_obcina(extracted_records)

                # --- NOVO: Samodejno poišči župana in matično, če je lastnik občina ---
                for uid, o_data in st.session_state.bulk_owners.items():
                    o_name = o_data["owners"][0].get("ime_priimek", "") if o_data["owners"] else ""
                    if "občina" in o_name.lower() or "obcina" in o_name.lower():
                        with st.spinner(f"Pridobivam dodatne podatke za '{o_name}' iz kdo-vodi.si..."):
                            z, m = pridobi_podatke_obcine(o_name)
                            if z:
                                o_data["zupan"] = z
                                o_data["maticna"] = m
                # ---------------------------------------------------------------------

                st.success(
                    f"Razporejenih {len(extracted_records)} izpiskov med "
                    f"{len(st.session_state.bulk_owners)} lastniško(imi) enoto(ami) "
                    "(solastniki iste parcele so združeni v eno enoto)."
                )
                for w in undercount_warnings:
                    st.warning(w)
                for w in find_name_conflicts(extracted_records):
                    st.warning(w)
            except LLMExtractionError as e:
                st.error(str(e))

    else:
        st.info(f"Najprej naložite PDF izpiske zemljiške knjige zgoraj (do {MAX_BULK_PDFS} hkrati).")

    if st.session_state.bulk_owners:
        st.divider()
        pcol, kcol = st.columns(2, gap="large")

        with pcol:
            st.subheader("👤 Pregled po lastniku")

            def _unit_names(data):
                return ", ".join(o["ime_priimek"] for o in data["owners"])

            def _unit_address(data):
                addrs, seen = [], set()
                for o in data["owners"]:
                    a = (o.get("naslov") or "").strip()
                    if a and a not in seen:
                        seen.add(a)
                        addrs.append(a)
                return "; ".join(addrs)

            owner_items = list(st.session_state.bulk_owners.items())
            owner_labels = [
                _unit_names(data)
                + f"  ({len(dedupe_parcels(data['records']))} parcel)"
                for _, data in owner_items
            ]
            owner_sel = st.selectbox(
                "Izberi lastnika/lastnico", options=range(len(owner_items)),
                format_func=lambda i: owner_labels[i], key="bulk_owner_select",
            )
            _, owner_data = owner_items[owner_sel]
            with st.container(border=True):
                st.markdown(
                    f"**{_unit_names(owner_data)}**"
                    + (f", {_unit_address(owner_data)}" if _unit_address(owner_data) else "")
                )
                
                # --- NOVO: Prikaz župana in matične št. v pregledu ---
                if owner_data.get("zupan"):
                    st.markdown(
                        f"<div style='margin-bottom: 8px; font-size: 0.95rem; color: #1E4A7B;'>"
                        f"🤵 <b>Župan/ja:</b> {owner_data['zupan']} &nbsp;&nbsp;|&nbsp;&nbsp; 🏢 <b>Matična št.:</b> {owner_data['maticna']}"
                        f"</div>", 
                        unsafe_allow_html=True
                    )
                # -------------------------------------------------------
                
                for rec in owner_data["records"]:
                    st.markdown(
                        f"- **{rec.get('ime_priimek') or '_ni najdeno_'}** — "
                        f"**{rec.get('katastrska_obcina') or '_ni najdeno_'}**, "
                        f"parc. št. **{rec.get('parcelna_stevilka') or '_ni najdeno_'}**, "
                        f"delež **{rec.get('delez') or '_ni najdeno_'}**"
                    )
                    st.markdown(
                        f'<div class="field-caption">vir: {rec.get("file_name", "")}</div>',
                        unsafe_allow_html=True,
                    )

            if st.button(
                "📤 PRENESI V IZPOLNJEVANJE POGODBE", key="bulk_transfer_owner",
                type="primary", use_container_width=True,
            ):
                if owner_data:
                    parcels_for_fill = [
                        {
                            "ime_priimek": rec.get("ime_priimek", ""),
                            "naslov": rec.get("naslov", ""),
                            "katastrska_obcina": rec.get("katastrska_obcina", ""),
                            "parcelna_stevilka": rec.get("parcelna_stevilka", ""),
                            "delez": rec.get("delez", ""),
                        }
                        for rec in owner_data["records"]
                    ]
                    owner_label = _unit_names(owner_data) + (
                        f", {_unit_address(owner_data)}" if _unit_address(owner_data) else ""
                    )
                else:
                    parcels_for_fill = []
                    owner_label = "Ni podatkov"

                st.session_state.transferred_parcels = parcels_for_fill
                st.session_state.transferred_owner_label = owner_label
                st.session_state.transferred_zupan = owner_data.get("zupan")
                st.session_state.transferred_maticna = owner_data.get("maticna")
                st.session_state.switch_to_fill_tab = True

                template_msg = None
                if parcels_for_fill:
                    unique_owners_for_template = dedupe_owners(parcels_for_fill)
                    tmpl_filename, tmpl_label, gender_uncertain = pick_contract_template(
                        unique_owners_for_template
                    )
                    tmpl_bytes = _load_template_bytes(tmpl_filename)
                    if tmpl_bytes:
                        _activate_template(tmpl_bytes, file_id=f"auto::{tmpl_filename}")
                        st.session_state.auto_template_info = {
                            "file_id": f"auto::{tmpl_filename}",
                            "label": tmpl_label,
                            "gender_uncertain": gender_uncertain,
                        }
                        template_msg = f" Samodejno naložena predloga: **{tmpl_label}**."
                    else:
                        st.session_state.auto_template_info = None
                        template_msg = (
                            f" ⚠️ Predloge '{tmpl_filename}' ni bilo mogoče najti v mapi "
                            "templates/ - naložite jo ročno v zavihku 'Izpolnjevanje pogodbe'."
                        )

                st.success(
                    f"Podatki za **{st.session_state.transferred_owner_label}** "
                    f"({len(parcels_for_fill)} vnos(ov)) so preneseni."
                    + (template_msg or "")
                    + " Preklapljam na zavihek '📝 Izpolnjevanje pogodbe' ..."
                )
                st.rerun()

        with kcol:
            st.subheader("🗺️ Pregled po katastrski občini")
            ko_items = list(st.session_state.bulk_by_ko.items())
            ko_labels = [f"{ko} ({len(data['entries'])} vnos(ov))" for ko, data in ko_items]
            ko_sel = st.selectbox(
                "Izberi katastrsko občino", options=range(len(ko_items)),
                format_func=lambda i: ko_labels[i], key="bulk_ko_select",
            )
            ko_name, ko_data = ko_items[ko_sel]
            with st.container(border=True):
                st.markdown(f"**{ko_name}**")
                for e in ko_data["entries"]:
                    st.markdown(
                        f"- **{e['ime_priimek'] or '_ni najdeno_'}**"
                        + (f", {e['naslov']}" if e["naslov"] else "")
                        + f" — parc. št. **{e['parcelna_stevilka'] or '_ni najdeno_'}**, "
                        f"delež **{e['delez'] or '_ni najdeno_'}**"
                    )
                    st.markdown(f'<div class="field-caption">vir: {e["source_file"]}</div>', unsafe_allow_html=True)

        st.divider()
        clear_col, reset_col = st.columns(2)
        with clear_col:
            if st.button(
                "🗑️ Počisti naložene datoteke PDF", use_container_width=True, key="bulk_clear_pdfs",
                help="Počisti seznam naloženih PDF izpiskov (lahko jih naložite znova) - "
                     "že izluščeni rezultati spodaj ostanejo.",
            ):
                st.session_state.bulk_pdf_records = []
                st.session_state.bulk_uploader_version += 1
                st.rerun()
        with reset_col:
            if st.button("🔄 Počisti skupinski pregled", use_container_width=True, key="bulk_reset"):
                st.session_state.bulk_pdf_records = []
                st.session_state.bulk_extracted = None
                st.session_state.bulk_owners = None
                st.session_state.bulk_by_ko = None
                st.session_state.bulk_uploader_version += 1
                st.rerun()
