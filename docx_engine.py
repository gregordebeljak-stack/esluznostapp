"""
docx_engine.py
--------------
Razčlenjevanje .docx predlog: najde vsa besedila, obarvana rdeče, jih združi
v logična vnosna polja (Field), omogoča urejanje vrednosti in izvoz nazaj
v .docx z rdečo barvo pretvorjeno v črno, ob ohranjanju originalne oblike.

Ker Streamlit ob vsaki interakciji znova požene skripto, celoten `Document`
objekt hranimo v st.session_state in polja urejamo *na živo* neposredno na
run-objektih iz python-docx (ne s tekstovnim iskanjem/zamenjavo), s čimer je
100-odstotno ohranjena oblika (pisava, velikost, tabele, odmiki, ...).
"""

from __future__ import annotations

import io
import uuid
from dataclasses import dataclass, field as dc_field
from typing import Dict, List, Optional, Tuple

import docx
from docx.document import Document as DocxDocument
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Barve, ki jih obravnavamo kot "rdeč placeholder". Uradna zahteva je #FF0000,
# a v realnih predlogah (preverjeno na primeru uporabnika) se pogosto
# uporablja tudi temnejša odtenka #C00000 ("Dark Red" v Wordovi paleti), zato
# ju privzeto zaznavamo oba. Uporabnik lahko seznam po potrebi razširi/skrči
# v UI-ju nadzorne plošče.
DEFAULT_RED_HEX = ["FF0000", "C00000"]

BLACK = RGBColor(0x00, 0x00, 0x00)


@dataclass
class FieldLocation:
    """Ena 'skupina' zaporednih rdečih run-ov na enem mestu v dokumentu."""
    runs: List[Run]
    paragraph: Paragraph
    context_before: str = ""   # besedilo tik pred skupino (za oznako polja)
    context_after: str = ""    # besedilo tik za skupino
    table_index: Optional[int] = None
    row_index: Optional[int] = None
    col_index: Optional[int] = None
    column_header: str = ""    # besedilo v glavi stolpca (vrstica 0), če gre za tabelo


@dataclass
class Field:
    id: str
    label: str
    value: str
    category: str
    is_blank_placeholder: bool
    locations: List[FieldLocation] = dc_field(default_factory=list)

    @property
    def context_hint(self) -> str:
        """Združeno besedilo konteksta vseh pojavitev - uporablja se za
        prepoznavanje vloge polja (ime, naslov, parcela, delež, k.o. ...)."""
        parts = [self.label]
        for loc in self.locations:
            parts.extend([loc.context_before, loc.context_after, loc.column_header])
        return " | ".join(p for p in parts if p)

    def to_ui_dict(self) -> dict:
        return {
            "id": self.id,
            "label": self.label,
            "value": self.value,
            "category": self.category,
            "is_blank_placeholder": self.is_blank_placeholder,
            "n_occurrences": len(self.locations),
        }


def _run_color_hex(run: Run) -> Optional[str]:
    try:
        color = run.font.color
        if color is not None and color.type is not None and color.rgb is not None:
            return str(color.rgb)
    except Exception:
        return None
    return None


def _is_target_red(run: Run, red_hex: List[str]) -> bool:
    hexval = _run_color_hex(run)
    return hexval is not None and hexval.upper() in {h.upper() for h in red_hex}


def _group_runs_in_paragraph(
    paragraph: Paragraph,
    red_hex: List[str],
    table_index: Optional[int] = None,
    row_index: Optional[int] = None,
    col_index: Optional[int] = None,
    column_header: str = "",
) -> List[FieldLocation]:
    """Poišče zaporedne rdeče run-e znotraj enega odstavka in jih združi."""
    groups: List[FieldLocation] = []
    runs = paragraph.runs
    i = 0
    n = len(runs)
    while i < n:
        if _is_target_red(runs[i], red_hex):
            start = i
            while i < n and _is_target_red(runs[i], red_hex):
                i += 1
            group_runs = runs[start:i]
            before = "".join(r.text for r in runs[:start])[-60:]
            after = "".join(r.text for r in runs[i:])[:60]
            groups.append(
                FieldLocation(
                    runs=group_runs,
                    paragraph=paragraph,
                    context_before=before,
                    context_after=after,
                    table_index=table_index,
                    row_index=row_index,
                    col_index=col_index,
                    column_header=column_header,
                )
            )
        else:
            i += 1
    return groups


def _categorize(label_text: str, context: str) -> str:
    text = f"{label_text} {context}".lower()
    if any(k in text for k in ["emšo", "ime", "priimek", "imetnik", "lastnik", "naslov"]):
        return "Lastnik / imetnik"
    if any(k in text for k in ["parc", "k.o.", "katastrsk", "nepremičnin", "id znak", "delež", "delez"]):
        return "Nepremičnina (parcela)"
    if any(k in text for k in ["breme", "služnost", "zaznamb"]):
        return "Bremena / služnosti"
    if any(k in text for k in ["datum", "sklenj", "dne "]):
        return "Datumi"
    return "Besedilo pogodbe"


def _make_label(loc: FieldLocation, fallback_text: str) -> str:
    before = loc.context_before.strip()
    after = loc.context_after.strip()
    if loc.table_index is not None:
        prefix = f"Tabela {loc.table_index + 1}"
        if before:
            return f"{prefix}: … {before[-30:]}"
        if after:
            return f"{prefix}: {after[:30]} …"
        return f"{prefix}, vrstica {loc.row_index + 1}, stolpec {loc.col_index + 1}"
    if before:
        return f"… {before[-40:]}"
    if fallback_text.strip():
        return fallback_text.strip()[:50]
    if after:
        return f"{after[:40]} …"
    return "Prazno polje (ročni vnos)"


def parse_docx_fields(doc: DocxDocument, red_hex: List[str] = None) -> Dict[str, Field]:
    """Preišče celoten dokument (odstavki + tabele) in vrne slovar polj."""
    if red_hex is None:
        red_hex = DEFAULT_RED_HEX

    fields_by_text: Dict[str, Field] = {}   # dedup po unikatnem besedilu (ne-prazna polja)
    fields_ordered: Dict[str, Field] = {}   # končni izhod, ohranja vrstni red

    def register_group(loc: FieldLocation):
        text = "".join(r.text for r in loc.runs)
        label = _make_label(loc, text)
        category = _categorize(label, loc.context_before + " " + loc.context_after + " " + loc.column_header)

        if text.strip() == "":
            # Prazno rdeče polje = ročno izpolnjevanje (npr. številka EMŠO).
            fid = f"field_{uuid.uuid4().hex[:8]}"
            f = Field(
                id=fid,
                label=label,
                value="",
                category=category,
                is_blank_placeholder=True,
                locations=[loc],
            )
            fields_ordered[fid] = f
        else:
            key = text.strip()
            if key in fields_by_text:
                fields_by_text[key].locations.append(loc)
            else:
                fid = f"field_{uuid.uuid4().hex[:8]}"
                f = Field(
                    id=fid,
                    label=label,
                    value=text,
                    category=category,
                    is_blank_placeholder=False,
                    locations=[loc],
                )
                fields_by_text[key] = f
                fields_ordered[fid] = f

    # 1) Odstavki na najvišji ravni (izven tabel)
    for p in doc.paragraphs:
        for loc in _group_runs_in_paragraph(p, red_hex):
            register_group(loc)

    # 2) Tabele (rekurzivno - celice lahko vsebujejo gnezdene tabele)
    def walk_table(table: Table, t_index: int):
        # Za vsak stolpec zberemo besedilo VSEH vrstic NAD trenutno vrstico
        # (ne le vrstice 0) - nekatere tabele imajo nad pravo vrstico glav
        # (npr. "parc. št. | k.o. | delež") še spojeno naslovno vrstico
        # (npr. "Miran Vajda" čez celo širino), zato "prava" glava ni nujno v vrstici 0.
        n_cols = len(table.columns) if table.columns else 0

        def header_hint(r_idx: int, c_idx: int) -> str:
            texts = []
            for above_r in range(0, r_idx):
                try:
                    t = table.rows[above_r].cells[c_idx].text.strip()
                except IndexError:
                    continue
                if t and t not in texts:
                    texts.append(t)
            return " | ".join(texts)

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                col_header = header_hint(r_idx, c_idx) if r_idx > 0 else ""
                for p in cell.paragraphs:
                    for loc in _group_runs_in_paragraph(
                        p, red_hex, table_index=t_index, row_index=r_idx, col_index=c_idx,
                        column_header=col_header,
                    ):
                        register_group(loc)
                for nested in cell.tables:
                    walk_table(nested, t_index)

    for t_idx, table in enumerate(doc.tables):
        walk_table(table, t_idx)

    return fields_ordered


def apply_field_value(f: Field, new_value: str):
    """Zapiše novo vrednost polja v VSE njegove pojavitve in obarva črno."""
    f.value = new_value
    for loc in f.locations:
        if not loc.runs:
            continue
        loc.runs[0].text = new_value
        loc.runs[0].font.color.rgb = BLACK
        for extra_run in loc.runs[1:]:
            extra_run.text = ""
            try:
                extra_run.font.color.rgb = BLACK
            except Exception:
                pass


def load_docx(file_bytes: bytes) -> DocxDocument:
    return docx.Document(io.BytesIO(file_bytes))


def save_docx(doc: DocxDocument) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
